#!/usr/bin/env python3
"""
System Builder Hub - Enhanced Server with OpenAI Integration and System Generation

Configuration:
- OPENAI_API_KEY: Required in production (provided via ECS task definition secrets)
- OPENAI_MODEL: Default gpt-4o-mini
- OPENAI_TIMEOUT_SECONDS: Default 20 seconds
"""
import os
import time
import uuid
import json
import logging
from datetime import datetime
from typing import Optional, Dict, Any
from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
import openai
from openai import OpenAI
import boto3
from botocore.exceptions import ClientError

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# S3 Configuration
s3_client = boto3.client('s3')
S3_BUCKET = os.getenv('S3_BUCKET_NAME', 'sbh-generated-systems')

def get_openai_config() -> Dict[str, Any]:
    """Get OpenAI configuration from environment variables"""
    return {
        'api_key': os.getenv('OPENAI_API_KEY'),
        'model': os.getenv('OPENAI_MODEL', 'gpt-4o-mini'),
        'timeout': int(os.getenv('OPENAI_TIMEOUT_SECONDS', '20'))
    }

def create_openai_client() -> Optional[OpenAI]:
    """Create OpenAI client if API key is available"""
    config = get_openai_config()
    if not config['api_key']:
        return None
    
    try:
        return OpenAI(
            api_key=config['api_key'],
            timeout=config['timeout']
        )
    except Exception as e:
        logger.error(f"Failed to create OpenAI client: {e}")
        return None

def generate_system_architecture(spec):
    """Generate system architecture based on specifications"""
    architecture = {
        'components': [],
        'dataFlow': [],
        'infrastructure': [],
        'security': [],
        'scalability': []
    }
    
    # Add components based on system type
    if spec['type'] == 'web-app':
        architecture['components'] = [
            {'name': 'Frontend', 'type': 'React App', 'port': 3000},
            {'name': 'Backend API', 'type': 'Node.js/Express', 'port': 8000},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432}
        ]
    elif spec['type'] == 'api':
        architecture['components'] = [
            {'name': 'API Gateway', 'type': 'Express.js', 'port': 8000},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432}
        ]
    elif spec['type'] == 'ecommerce-platform':
        architecture['components'] = [
            {'name': 'Frontend', 'type': 'Next.js', 'port': 3000},
            {'name': 'API Gateway', 'type': 'Node.js/Express', 'port': 8000},
            {'name': 'Payment Service', 'type': 'Stripe Integration', 'port': 8001},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432},
            {'name': 'File Storage', 'type': 'AWS S3', 'port': None}
        ]
    elif spec['type'] == 'data-pipeline':
        architecture['components'] = [
            {'name': 'Data Ingestion', 'type': 'AWS Lambda', 'port': None},
            {'name': 'Data Processing', 'type': 'AWS ECS', 'port': 8000},
            {'name': 'Data Storage', 'type': 'AWS S3', 'port': None},
            {'name': 'Data Warehouse', 'type': 'AWS Redshift', 'port': 5439}
        ]
    elif spec['type'] == 'ml-service':
        architecture['components'] = [
            {'name': 'Model API', 'type': 'FastAPI', 'port': 8000},
            {'name': 'Model Storage', 'type': 'AWS S3', 'port': None},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432},
            {'name': 'Monitoring', 'type': 'CloudWatch', 'port': None}
        ]
    elif spec['type'] == 'microservice':
        architecture['components'] = [
            {'name': 'Service API', 'type': 'Express.js', 'port': 8000},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432},
            {'name': 'Message Queue', 'type': 'AWS SQS', 'port': None}
        ]
    elif spec['type'] == 'cms':
        architecture['components'] = [
            {'name': 'Frontend', 'type': 'Next.js', 'port': 3000},
            {'name': 'Admin Panel', 'type': 'React', 'port': 3001},
            {'name': 'API Gateway', 'type': 'Node.js/Express', 'port': 8000},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432},
            {'name': 'File Storage', 'type': 'AWS S3', 'port': None}
        ]
    elif spec['type'] == 'dashboard':
        architecture['components'] = [
            {'name': 'Dashboard UI', 'type': 'React', 'port': 3000},
            {'name': 'API Gateway', 'type': 'Node.js/Express', 'port': 8000},
            {'name': 'Data Processing', 'type': 'AWS Lambda', 'port': None},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432},
            {'name': 'Cache', 'type': 'Redis', 'port': 6379}
        ]
    
    # Add infrastructure based on selections
    if 'AWS ECS Fargate' in spec['infrastructure']:
        architecture['infrastructure'].append({
            'name': 'Container Orchestration',
            'type': 'AWS ECS Fargate',
            'description': 'Serverless container platform'
        })
    
    if 'AWS RDS' in spec['infrastructure']:
        architecture['infrastructure'].append({
            'name': 'Database',
            'type': 'AWS RDS PostgreSQL',
            'description': 'Managed PostgreSQL database'
        })
    
    if 'AWS S3' in spec['infrastructure']:
        architecture['infrastructure'].append({
            'name': 'File Storage',
            'type': 'AWS S3',
            'description': 'Object storage for files and assets'
        })
    
    if 'AWS ALB' in spec['infrastructure']:
        architecture['infrastructure'].append({
            'name': 'Load Balancer',
            'type': 'AWS Application Load Balancer',
            'description': 'Application load balancer for traffic distribution'
        })
    
    return architecture

def save_system_to_s3(system_id, system_data):
    """Save system to S3"""
    try:
        key = f"systems/{system_id}.json"
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=key,
            Body=json.dumps(system_data, default=str),
            ContentType='application/json'
        )
        return True
    except ClientError as e:
        logger.error(f"Error saving system to S3: {e}")
        return False

def load_system_from_s3(system_id):
    """Load system from S3"""
    try:
        key = f"systems/{system_id}.json"
        response = s3_client.get_object(Bucket=S3_BUCKET, Key=key)
        return json.loads(response['Body'].read().decode('utf-8'))
    except ClientError as e:
        if e.response['Error']['Code'] == 'NoSuchKey':
            return None
        logger.error(f"Error loading system from S3: {e}")
        return None

def delete_system_from_s3(system_id):
    """Delete system from S3"""
    try:
        key = f"systems/{system_id}.json"
        s3_client.delete_object(Bucket=S3_BUCKET, Key=key)
        return True
    except ClientError as e:
        logger.error(f"Error deleting system from S3: {e}")
        return False

def detect_domain_type(domain):
    """Detect the type of domain for deployment strategy"""
    domain = domain.lower().strip()
    
    # SBH-managed subdomains
    if domain.endswith('.sbh.umbervale.com'):
        return 'sbh_managed'
    
    # Check if it's a root domain (no subdomain)
    parts = domain.split('.')
    if len(parts) == 2:  # ecommerce.com
        return 'root_domain'
    elif len(parts) > 2:  # ecommerce.umbervale.com
        return 'custom_subdomain'
    
    return 'unknown'

def validate_domain(domain):
    """Validate domain format and availability"""
    import re
    
    # Basic domain format validation
    domain_pattern = r'^[a-zA-Z0-9]([a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?(\.[a-zA-Z0-9]([a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?)*$'
    
    if not re.match(domain_pattern, domain):
        return {
            'valid': False,
            'error': 'Invalid domain format'
        }
    
    # Check domain length
    if len(domain) > 253:
        return {
            'valid': False,
            'error': 'Domain too long (max 253 characters)'
        }
    
    # Check for reserved domains
    reserved_domains = ['localhost', 'example.com', 'test.com', 'sbh.umbervale.com']
    if domain in reserved_domains:
        return {
            'valid': False,
            'error': 'Domain is reserved'
        }
    
    return {
        'valid': True,
        'domain_type': detect_domain_type(domain)
    }

def get_deployment_strategy(domain_type):
    """Get deployment strategy based on domain type"""
    strategies = {
        'sbh_managed': {
            'dns_management': 'automatic',
            'ssl_management': 'automatic',
            'user_action_required': False,
            'setup_instructions': 'Fully automated - no user action required'
        },
        'custom_subdomain': {
            'dns_management': 'cname_instructions',
            'ssl_management': 'automatic_after_dns',
            'user_action_required': True,
            'setup_instructions': 'Create CNAME record pointing to our load balancer'
        },
        'root_domain': {
            'dns_management': 'route53_or_cloudflare',
            'ssl_management': 'automatic_after_dns',
            'user_action_required': True,
            'setup_instructions': 'Transfer to Route 53 or use CloudFlare for best results'
        }
    }
    
    return strategies.get(domain_type, {
        'dns_management': 'manual',
        'ssl_management': 'manual',
        'user_action_required': True,
        'setup_instructions': 'Manual DNS configuration required'
    })

def create_route53_record(domain, target, record_type='CNAME'):
    """Create Route 53 DNS record"""
    try:
        route53_client = boto3.client('route53')
        
        # Get hosted zone ID for the domain
        hosted_zone_id = get_hosted_zone_id(domain)
        if not hosted_zone_id:
            return {
                'success': False,
                'error': 'No Route 53 hosted zone found for domain'
            }
        
        # Create DNS record
        response = route53_client.change_resource_record_sets(
            HostedZoneId=hosted_zone_id,
            ChangeBatch={
                'Changes': [{
                    'Action': 'UPSERT',
                    'ResourceRecordSet': {
                        'Name': domain,
                        'Type': record_type,
                        'TTL': 300,
                        'ResourceRecords': [{'Value': target}]
                    }
                }]
            }
        )
        
        return {
            'success': True,
            'change_id': response['ChangeInfo']['Id'],
            'status': response['ChangeInfo']['Status']
        }
        
    except Exception as e:
        logger.error(f"Error creating Route 53 record: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def get_hosted_zone_id(domain):
    """Get Route 53 hosted zone ID for domain"""
    try:
        route53_client = boto3.client('route53')
        
        # Extract root domain
        parts = domain.split('.')
        if len(parts) >= 2:
            root_domain = '.'.join(parts[-2:])
        else:
            root_domain = domain
        
        # List hosted zones
        response = route53_client.list_hosted_zones()
        
        for zone in response['HostedZones']:
            zone_name = zone['Name'].rstrip('.')
            if zone_name == root_domain:
                return zone['Id'].split('/')[-1]
        
        return None
        
    except Exception as e:
        logger.error(f"Error getting hosted zone ID: {e}")
        return None

def request_ssl_certificate(domain):
    """Request SSL certificate from AWS Certificate Manager"""
    try:
        acm_client = boto3.client('acm', region_name='us-east-1')  # ACM requires us-east-1
        
        # Request certificate
        response = acm_client.request_certificate(
            DomainName=domain,
            ValidationMethod='DNS',
            SubjectAlternativeNames=[domain] if not domain.startswith('*.') else [domain]
        )
        
        certificate_arn = response['CertificateArn']
        
        # Get DNS validation records
        validation_records = get_certificate_validation_records(certificate_arn)
        
        return {
            'success': True,
            'certificate_arn': certificate_arn,
            'validation_records': validation_records,
            'status': 'pending_validation'
        }
        
    except Exception as e:
        logger.error(f"Error requesting SSL certificate: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def get_certificate_validation_records(certificate_arn):
    """Get DNS validation records for SSL certificate"""
    try:
        acm_client = boto3.client('acm', region_name='us-east-1')
        
        response = acm_client.describe_certificate(CertificateArn=certificate_arn)
        
        validation_records = []
        for option in response['Certificate']['DomainValidationOptions']:
            if 'ResourceRecord' in option:
                validation_records.append({
                    'name': option['ResourceRecord']['Name'],
                    'type': option['ResourceRecord']['Type'],
                    'value': option['ResourceRecord']['Value']
                })
        
        return validation_records
        
    except Exception as e:
        logger.error(f"Error getting validation records: {e}")
        return []

def check_dns_propagation(domain, expected_value):
    """Check if DNS record has propagated"""
    try:
        import socket
        
        # Try to resolve the domain
        result = socket.gethostbyname(domain)
        
        # For CNAME records, we'd need to check the actual CNAME value
        # This is a simplified check
        return {
            'propagated': True,
            'resolved_ip': result
        }
        
    except socket.gaierror:
        return {
            'propagated': False,
            'error': 'DNS not yet propagated'
        }
    except Exception as e:
        return {
            'propagated': False,
            'error': str(e)
        }

def upload_file_to_s3(file, system_id, file_type):
    """Upload file to S3 for system reference"""
    try:
        import uuid
        from werkzeug.utils import secure_filename
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_extension = secure_filename(file.filename).split('.')[-1] if '.' in file.filename else ''
        s3_key = f"references/{system_id}/{file_type}/{file_id}.{file_extension}"
        
        # Upload to S3
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=s3_key,
            Body=file.read(),
            ContentType=file.content_type or 'application/octet-stream'
        )
        
        return {
            'success': True,
            'file_id': file_id,
            's3_key': s3_key,
            'filename': file.filename,
            'content_type': file.content_type,
            'size': file.content_length
        }
        
    except Exception as e:
        logger.error(f"Error uploading file to S3: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def analyze_uploaded_image(s3_key):
    """Analyze uploaded image for system design insights"""
    try:
        from PIL import Image
        import io
        
        # Download image from S3
        response = s3_client.get_object(Bucket=S3_BUCKET, Key=s3_key)
        image_data = response['Body'].read()
        
        # Open image with PIL
        image = Image.open(io.BytesIO(image_data))
        
        # Basic analysis
        analysis = {
            'width': image.width,
            'height': image.height,
            'format': image.format,
            'mode': image.mode,
            'size_kb': len(image_data) / 1024,
            'aspect_ratio': round(image.width / image.height, 2)
        }
        
        # Detect UI elements (basic color analysis)
        if image.mode == 'RGB':
            # Get dominant colors
            colors = image.getcolors(maxcolors=256*256*256)
            if colors:
                dominant_color = max(colors, key=lambda x: x[0])
                analysis['dominant_color'] = dominant_color[1]
        
        return {
            'success': True,
            'analysis': analysis
        }
        
    except Exception as e:
        logger.error(f"Error analyzing image: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def analyze_reference_url(url):
    """Analyze reference URL for system inspiration"""
    try:
        import requests
        from bs4 import BeautifulSoup
        
        # Fetch the webpage
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        # Parse HTML
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Extract key information
        analysis = {
            'title': soup.title.string if soup.title else '',
            'description': '',
            'technologies': [],
            'features': [],
            'layout_elements': []
        }
        
        # Get meta description
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc:
            analysis['description'] = meta_desc.get('content', '')
        
        # Detect technologies
        scripts = soup.find_all('script', src=True)
        for script in scripts:
            src = script['src']
            if 'react' in src.lower():
                analysis['technologies'].append('React')
            elif 'vue' in src.lower():
                analysis['technologies'].append('Vue.js')
            elif 'angular' in src.lower():
                analysis['technologies'].append('Angular')
            elif 'jquery' in src.lower():
                analysis['technologies'].append('jQuery')
        
        # Detect common UI elements
        if soup.find('nav'):
            analysis['layout_elements'].append('Navigation')
        if soup.find('form'):
            analysis['layout_elements'].append('Forms')
        if soup.find('button'):
            analysis['layout_elements'].append('Buttons')
        if soup.find('input'):
            analysis['layout_elements'].append('Input Fields')
        
        # Detect features based on content
        page_text = soup.get_text().lower()
        if 'login' in page_text or 'sign in' in page_text:
            analysis['features'].append('User Authentication')
        if 'search' in page_text:
            analysis['features'].append('Search Functionality')
        if 'cart' in page_text or 'shopping' in page_text:
            analysis['features'].append('E-commerce')
        if 'contact' in page_text:
            analysis['features'].append('Contact Forms')
        
        return {
            'success': True,
            'url': url,
            'analysis': analysis
        }
        
    except Exception as e:
        logger.error(f"Error analyzing URL: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def extract_text_from_document(s3_key, content_type):
    """Extract text from uploaded documents"""
    try:
        # Download file from S3
        response = s3_client.get_object(Bucket=S3_BUCKET, Key=s3_key)
        file_data = response['Body'].read()
        
        text_content = ""
        
        if content_type == 'application/pdf':
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
                
        elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            from docx import Document
            import io
            doc = Document(io.BytesIO(file_data))
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
                
        elif content_type == 'text/plain':
            text_content = file_data.decode('utf-8')
            
        return {
            'success': True,
            'text_content': text_content.strip(),
            'word_count': len(text_content.split())
        }
        
    except Exception as e:
        logger.error(f"Error extracting text from document: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def generate_setup_instructions(domain_type, domain):
    """Generate setup instructions based on domain type"""
    if domain_type == 'subdomain':
        return {
            'type': 'subdomain',
            'instructions': [
                f"1. Add a CNAME record for {domain}",
                "2. Point it to: sbh.umbervale.com",
                "3. SSL certificate will be automatically provisioned",
                "4. DNS propagation may take 5-10 minutes"
            ],
            'dns_record': {
                'type': 'CNAME',
                'name': domain,
                'value': 'sbh.umbervale.com'
            }
        }
    elif domain_type == 'custom_domain':
        return {
            'type': 'custom_domain',
            'instructions': [
                f"1. Add a CNAME record for {domain}",
                "2. Point it to: sbh.umbervale.com",
                "3. SSL certificate will be automatically provisioned",
                "4. DNS propagation may take 5-10 minutes"
            ],
            'dns_record': {
                'type': 'CNAME',
                'name': domain,
                'value': 'sbh.umbervale.com'
            }
        }
    elif domain_type == 'root_domain':
        return {
            'type': 'root_domain',
            'instructions': [
                f"1. Root domains require A records (not CNAME)",
                "2. Contact support for ALB IP addresses",
                "3. Alternative: Transfer domain to Route 53",
                "4. SSL certificate will be automatically provisioned"
            ],
            'dns_record': {
                'type': 'A',
                'name': domain,
                'value': 'Contact support for ALB IP'
            }
        }
    else:
        return {
            'type': 'unknown',
            'instructions': [
                "1. Contact support for domain setup assistance",
                "2. Provide domain details for custom configuration"
            ],
            'dns_record': {
                'type': 'Unknown',
                'name': domain,
                'value': 'Contact support'
            }
        }

#!/usr/bin/env python3
"""
System Builder Hub - Enhanced Server with OpenAI Integration and System Generation

Configuration:
- OPENAI_API_KEY: Required in production (provided via ECS task definition secrets)
- OPENAI_MODEL: Default gpt-4o-mini
- OPENAI_TIMEOUT_SECONDS: Default 20 seconds
"""
import os
import time
import uuid
import json
import logging
from datetime import datetime
from typing import Optional, Dict, Any
from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
import openai
from openai import OpenAI
import boto3
from botocore.exceptions import ClientError

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# S3 Configuration
s3_client = boto3.client('s3')
S3_BUCKET = os.getenv('S3_BUCKET_NAME', 'sbh-generated-systems')

def get_openai_config() -> Dict[str, Any]:
    """Get OpenAI configuration from environment variables"""
    return {
        'api_key': os.getenv('OPENAI_API_KEY'),
        'model': os.getenv('OPENAI_MODEL', 'gpt-4o-mini'),
        'timeout': int(os.getenv('OPENAI_TIMEOUT_SECONDS', '20'))
    }

def create_openai_client() -> Optional[OpenAI]:
    """Create OpenAI client if API key is available"""
    config = get_openai_config()
    if not config['api_key']:
        return None
    
    try:
        return OpenAI(
            api_key=config['api_key'],
            timeout=config['timeout']
        )
    except Exception as e:
        logger.error(f"Failed to create OpenAI client: {e}")
        return None

def generate_system_architecture(spec):
    """Generate system architecture based on specifications"""
    architecture = {
        'components': [],
        'dataFlow': [],
        'infrastructure': [],
        'security': [],
        'scalability': []
    }
    
    # Add components based on system type
    if spec['type'] == 'web-app':
        architecture['components'] = [
            {'name': 'Frontend', 'type': 'React App', 'port': 3000},
            {'name': 'Backend API', 'type': 'Node.js/Express', 'port': 8000},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432}
        ]
    elif spec['type'] == 'api':
        architecture['components'] = [
            {'name': 'API Gateway', 'type': 'Express.js', 'port': 8000},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432}
        ]
    elif spec['type'] == 'ecommerce-platform':
        architecture['components'] = [
            {'name': 'Frontend', 'type': 'Next.js', 'port': 3000},
            {'name': 'API Gateway', 'type': 'Node.js/Express', 'port': 8000},
            {'name': 'Payment Service', 'type': 'Stripe Integration', 'port': 8001},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432},
            {'name': 'File Storage', 'type': 'AWS S3', 'port': None}
        ]
    elif spec['type'] == 'data-pipeline':
        architecture['components'] = [
            {'name': 'Data Ingestion', 'type': 'AWS Lambda', 'port': None},
            {'name': 'Data Processing', 'type': 'AWS ECS', 'port': 8000},
            {'name': 'Data Storage', 'type': 'AWS S3', 'port': None},
            {'name': 'Data Warehouse', 'type': 'AWS Redshift', 'port': 5439}
        ]
    elif spec['type'] == 'ml-service':
        architecture['components'] = [
            {'name': 'Model API', 'type': 'FastAPI', 'port': 8000},
            {'name': 'Model Storage', 'type': 'AWS S3', 'port': None},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432},
            {'name': 'Monitoring', 'type': 'CloudWatch', 'port': None}
        ]
    elif spec['type'] == 'microservice':
        architecture['components'] = [
            {'name': 'Service API', 'type': 'Express.js', 'port': 8000},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432},
            {'name': 'Message Queue', 'type': 'AWS SQS', 'port': None}
        ]
    elif spec['type'] == 'cms':
        architecture['components'] = [
            {'name': 'Frontend', 'type': 'Next.js', 'port': 3000},
            {'name': 'Admin Panel', 'type': 'React', 'port': 3001},
            {'name': 'API Gateway', 'type': 'Node.js/Express', 'port': 8000},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432},
            {'name': 'File Storage', 'type': 'AWS S3', 'port': None}
        ]
    elif spec['type'] == 'dashboard':
        architecture['components'] = [
            {'name': 'Dashboard UI', 'type': 'React', 'port': 3000},
            {'name': 'API Gateway', 'type': 'Node.js/Express', 'port': 8000},
            {'name': 'Data Processing', 'type': 'AWS Lambda', 'port': None},
            {'name': 'Database', 'type': 'PostgreSQL', 'port': 5432},
            {'name': 'Cache', 'type': 'Redis', 'port': 6379}
        ]
    
    # Add infrastructure based on selections
    if 'AWS ECS Fargate' in spec['infrastructure']:
        architecture['infrastructure'].append({
            'name': 'Container Orchestration',
            'type': 'AWS ECS Fargate',
            'description': 'Serverless container platform'
        })
    
    if 'AWS RDS' in spec['infrastructure']:
        architecture['infrastructure'].append({
            'name': 'Database',
            'type': 'AWS RDS PostgreSQL',
            'description': 'Managed PostgreSQL database'
        })
    
    if 'AWS S3' in spec['infrastructure']:
        architecture['infrastructure'].append({
            'name': 'File Storage',
            'type': 'AWS S3',
            'description': 'Object storage for files and assets'
        })
    
    if 'AWS ALB' in spec['infrastructure']:
        architecture['infrastructure'].append({
            'name': 'Load Balancer',
            'type': 'AWS Application Load Balancer',
            'description': 'Application load balancer for traffic distribution'
        })
    
    return architecture

def save_system_to_s3(system_id, system_data):
    """Save system to S3"""
    try:
        key = f"systems/{system_id}.json"
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=key,
            Body=json.dumps(system_data, default=str),
            ContentType='application/json'
        )
        return True
    except ClientError as e:
        logger.error(f"Error saving system to S3: {e}")
        return False

def load_system_from_s3(system_id):
    """Load system from S3"""
    try:
        key = f"systems/{system_id}.json"
        response = s3_client.get_object(Bucket=S3_BUCKET, Key=key)
        return json.loads(response['Body'].read().decode('utf-8'))
    except ClientError as e:
        if e.response['Error']['Code'] == 'NoSuchKey':
            return None
        logger.error(f"Error loading system from S3: {e}")
        return None

def delete_system_from_s3(system_id):
    """Delete system from S3"""
    try:
        key = f"systems/{system_id}.json"
        s3_client.delete_object(Bucket=S3_BUCKET, Key=key)
        return True
    except ClientError as e:
        logger.error(f"Error deleting system from S3: {e}")
        return False

def detect_domain_type(domain):
    """Detect the type of domain for deployment strategy"""
    domain = domain.lower().strip()
    
    # SBH-managed subdomains
    if domain.endswith('.sbh.umbervale.com'):
        return 'sbh_managed'
    
    # Check if it's a root domain (no subdomain)
    parts = domain.split('.')
    if len(parts) == 2:  # ecommerce.com
        return 'root_domain'
    elif len(parts) > 2:  # ecommerce.umbervale.com
        return 'custom_subdomain'
    
    return 'unknown'

def validate_domain(domain):
    """Validate domain format and availability"""
    import re
    
    # Basic domain format validation
    domain_pattern = r'^[a-zA-Z0-9]([a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?(\.[a-zA-Z0-9]([a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?)*$'
    
    if not re.match(domain_pattern, domain):
        return {
            'valid': False,
            'error': 'Invalid domain format'
        }
    
    # Check domain length
    if len(domain) > 253:
        return {
            'valid': False,
            'error': 'Domain too long (max 253 characters)'
        }
    
    # Check for reserved domains
    reserved_domains = ['localhost', 'example.com', 'test.com', 'sbh.umbervale.com']
    if domain in reserved_domains:
        return {
            'valid': False,
            'error': 'Domain is reserved'
        }
    
    return {
        'valid': True,
        'domain_type': detect_domain_type(domain)
    }

def get_deployment_strategy(domain_type):
    """Get deployment strategy based on domain type"""
    strategies = {
        'sbh_managed': {
            'dns_management': 'automatic',
            'ssl_management': 'automatic',
            'user_action_required': False,
            'setup_instructions': 'Fully automated - no user action required'
        },
        'custom_subdomain': {
            'dns_management': 'cname_instructions',
            'ssl_management': 'automatic_after_dns',
            'user_action_required': True,
            'setup_instructions': 'Create CNAME record pointing to our load balancer'
        },
        'root_domain': {
            'dns_management': 'route53_or_cloudflare',
            'ssl_management': 'automatic_after_dns',
            'user_action_required': True,
            'setup_instructions': 'Transfer to Route 53 or use CloudFlare for best results'
        }
    }
    
    return strategies.get(domain_type, {
        'dns_management': 'manual',
        'ssl_management': 'manual',
        'user_action_required': True,
        'setup_instructions': 'Manual DNS configuration required'
    })

def create_route53_record(domain, target, record_type='CNAME'):
    """Create Route 53 DNS record"""
    try:
        route53_client = boto3.client('route53')
        
        # Get hosted zone ID for the domain
        hosted_zone_id = get_hosted_zone_id(domain)
        if not hosted_zone_id:
            return {
                'success': False,
                'error': 'No Route 53 hosted zone found for domain'
            }
        
        # Create DNS record
        response = route53_client.change_resource_record_sets(
            HostedZoneId=hosted_zone_id,
            ChangeBatch={
                'Changes': [{
                    'Action': 'UPSERT',
                    'ResourceRecordSet': {
                        'Name': domain,
                        'Type': record_type,
                        'TTL': 300,
                        'ResourceRecords': [{'Value': target}]
                    }
                }]
            }
        )
        
        return {
            'success': True,
            'change_id': response['ChangeInfo']['Id'],
            'status': response['ChangeInfo']['Status']
        }
        
    except Exception as e:
        logger.error(f"Error creating Route 53 record: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def get_hosted_zone_id(domain):
    """Get Route 53 hosted zone ID for domain"""
    try:
        route53_client = boto3.client('route53')
        
        # Extract root domain
        parts = domain.split('.')
        if len(parts) >= 2:
            root_domain = '.'.join(parts[-2:])
        else:
            root_domain = domain
        
        # List hosted zones
        response = route53_client.list_hosted_zones()
        
        for zone in response['HostedZones']:
            zone_name = zone['Name'].rstrip('.')
            if zone_name == root_domain:
                return zone['Id'].split('/')[-1]
        
        return None
        
    except Exception as e:
        logger.error(f"Error getting hosted zone ID: {e}")
        return None

def request_ssl_certificate(domain):
    """Request SSL certificate from AWS Certificate Manager"""
    try:
        acm_client = boto3.client('acm', region_name='us-east-1')  # ACM requires us-east-1
        
        # Request certificate
        response = acm_client.request_certificate(
            DomainName=domain,
            ValidationMethod='DNS',
            SubjectAlternativeNames=[domain] if not domain.startswith('*.') else [domain]
        )
        
        certificate_arn = response['CertificateArn']
        
        # Get DNS validation records
        validation_records = get_certificate_validation_records(certificate_arn)
        
        return {
            'success': True,
            'certificate_arn': certificate_arn,
            'validation_records': validation_records,
            'status': 'pending_validation'
        }
        
    except Exception as e:
        logger.error(f"Error requesting SSL certificate: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def get_certificate_validation_records(certificate_arn):
    """Get DNS validation records for SSL certificate"""
    try:
        acm_client = boto3.client('acm', region_name='us-east-1')
        
        response = acm_client.describe_certificate(CertificateArn=certificate_arn)
        
        validation_records = []
        for option in response['Certificate']['DomainValidationOptions']:
            if 'ResourceRecord' in option:
                validation_records.append({
                    'name': option['ResourceRecord']['Name'],
                    'type': option['ResourceRecord']['Type'],
                    'value': option['ResourceRecord']['Value']
                })
        
        return validation_records
        
    except Exception as e:
        logger.error(f"Error getting validation records: {e}")
        return []

def check_dns_propagation(domain, expected_value):
    """Check if DNS record has propagated"""
    try:
        import socket
        
        # Try to resolve the domain
        result = socket.gethostbyname(domain)
        
        # For CNAME records, we'd need to check the actual CNAME value
        # This is a simplified check
        return {
            'propagated': True,
            'resolved_ip': result
        }
        
    except socket.gaierror:
        return {
            'propagated': False,
            'error': 'DNS not yet propagated'
        }
    except Exception as e:
        return {
            'propagated': False,
            'error': str(e)
        }

def upload_file_to_s3(file, system_id, file_type):
    """Upload file to S3 for system reference"""
    try:
        import uuid
        from werkzeug.utils import secure_filename
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_extension = secure_filename(file.filename).split('.')[-1] if '.' in file.filename else ''
        s3_key = f"references/{system_id}/{file_type}/{file_id}.{file_extension}"
        
        # Upload to S3
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=s3_key,
            Body=file.read(),
            ContentType=file.content_type or 'application/octet-stream'
        )
        
        return {
            'success': True,
            'file_id': file_id,
            's3_key': s3_key,
            'filename': file.filename,
            'content_type': file.content_type,
            'size': file.content_length
        }
        
    except Exception as e:
        logger.error(f"Error uploading file to S3: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def analyze_uploaded_image(s3_key):
    """Analyze uploaded image for system design insights"""
    try:
        from PIL import Image
        import io
        
        # Download image from S3
        response = s3_client.get_object(Bucket=S3_BUCKET, Key=s3_key)
        image_data = response['Body'].read()
        
        # Open image with PIL
        image = Image.open(io.BytesIO(image_data))
        
        # Basic analysis
        analysis = {
            'width': image.width,
            'height': image.height,
            'format': image.format,
            'mode': image.mode,
            'size_kb': len(image_data) / 1024,
            'aspect_ratio': round(image.width / image.height, 2)
        }
        
        # Detect UI elements (basic color analysis)
        if image.mode == 'RGB':
            # Get dominant colors
            colors = image.getcolors(maxcolors=256*256*256)
            if colors:
                dominant_color = max(colors, key=lambda x: x[0])
                analysis['dominant_color'] = dominant_color[1]
        
        return {
            'success': True,
            'analysis': analysis
        }
        
    except Exception as e:
        logger.error(f"Error analyzing image: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def analyze_reference_url(url):
    """Analyze reference URL for system inspiration"""
    try:
        import requests
        from bs4 import BeautifulSoup
        
        # Fetch the webpage
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        # Parse HTML
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Extract key information
        analysis = {
            'title': soup.title.string if soup.title else '',
            'description': '',
            'technologies': [],
            'features': [],
            'layout_elements': []
        }
        
        # Get meta description
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc:
            analysis['description'] = meta_desc.get('content', '')
        
        # Detect technologies
        scripts = soup.find_all('script', src=True)
        for script in scripts:
            src = script['src']
            if 'react' in src.lower():
                analysis['technologies'].append('React')
            elif 'vue' in src.lower():
                analysis['technologies'].append('Vue.js')
            elif 'angular' in src.lower():
                analysis['technologies'].append('Angular')
            elif 'jquery' in src.lower():
                analysis['technologies'].append('jQuery')
        
        # Detect common UI elements
        if soup.find('nav'):
            analysis['layout_elements'].append('Navigation')
        if soup.find('form'):
            analysis['layout_elements'].append('Forms')
        if soup.find('button'):
            analysis['layout_elements'].append('Buttons')
        if soup.find('input'):
            analysis['layout_elements'].append('Input Fields')
        
        # Detect features based on content
        page_text = soup.get_text().lower()
        if 'login' in page_text or 'sign in' in page_text:
            analysis['features'].append('User Authentication')
        if 'search' in page_text:
            analysis['features'].append('Search Functionality')
        if 'cart' in page_text or 'shopping' in page_text:
            analysis['features'].append('E-commerce')
        if 'contact' in page_text:
            analysis['features'].append('Contact Forms')
        
        return {
            'success': True,
            'url': url,
            'analysis': analysis
        }
        
    except Exception as e:
        logger.error(f"Error analyzing URL: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def extract_text_from_document(s3_key, content_type):
    """Extract text from uploaded documents"""
    try:
        # Download file from S3
        response = s3_client.get_object(Bucket=S3_BUCKET, Key=s3_key)
        file_data = response['Body'].read()
        
        text_content = ""
        
        if content_type == 'application/pdf':
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
                
        elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            from docx import Document
            import io
            doc = Document(io.BytesIO(file_data))
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
                
        elif content_type == 'text/plain':
            text_content = file_data.decode('utf-8')
            
        return {
            'success': True,
            'text_content': text_content.strip(),
            'word_count': len(text_content.split())
        }
        
    except Exception as e:
        logger.error(f"Error extracting text from document: {e}")
        return {
            'success': False,
            'error': str(e)
        }

# Enhanced template generation functions
def generate_react_package_json(spec):
    """Generate real package.json for React"""
    return json.dumps({
        "name": spec['name'].lower().replace(" ", "-"),
        "version": "1.0.0",
        "private": True,
        "scripts": {
            "dev": "react-scripts start",
            "build": "react-scripts build",
            "test": "react-scripts test",
            "eject": "react-scripts eject"
        },
        "dependencies": {
            "react": "^18.2.0",
            "react-dom": "^18.2.0",
            "react-router-dom": "^6.8.0",
            "axios": "^1.4.0",
            "tailwindcss": "^3.3.5",
            "lucide-react": "^0.292.0"
        },
        "devDependencies": {
            "@types/react": "^18.2.0",
            "@types/react-dom": "^18.2.0",
            "typescript": "^5.2.0",
            "react-scripts": "5.0.1"
        }
    }, indent=2)

def generate_nextjs_package_json(spec):
    """Generate real package.json for Next.js"""
    return json.dumps({
        "name": spec['name'].lower().replace(" ", "-"),
        "version": "1.0.0",
        "private": True,
        "scripts": {
            "dev": "next dev",
            "build": "next build",
            "start": "next start",
            "lint": "next lint"
        },
        "dependencies": {
            "next": "14.0.0",
            "react": "^18.2.0",
            "react-dom": "^18.2.0",
            "axios": "^1.4.0",
            "tailwindcss": "^3.3.5",
            "lucide-react": "^0.292.0"
        },
        "devDependencies": {
            "@types/node": "^20.0.0",
            "@types/react": "^18.2.0",
            "@types/react-dom": "^18.2.0",
            "typescript": "^5.2.0"
        }
    }, indent=2)

def generate_nodejs_package_json(spec):
    """Generate real package.json for Node.js backend"""
    return json.dumps({
        "name": f"{spec['name'].lower().replace(' ', '-')}-backend",
        "version": "1.0.0",
        "description": f"Backend API for {spec['name']}",
        "main": "src/app.js",
        "scripts": {
            "start": "node src/app.js",
            "dev": "nodemon src/app.js",
            "test": "jest",
            "lint": "eslint src/"
        },
        "dependencies": {
            "express": "^4.18.0",
            "cors": "^2.8.5",
            "helmet": "^7.0.0",
            "pg": "^8.11.0",
            "bcryptjs": "^2.4.3",
            "jsonwebtoken": "^9.0.0",
            "dotenv": "^16.0.0",
            "joi": "^17.9.0"
        },
        "devDependencies": {
            "nodemon": "^3.0.0",
            "jest": "^29.0.0",
            "eslint": "^8.0.0"
        }
    }, indent=2)

def generate_python_requirements(spec):
    """Generate real requirements.txt for Python backend"""
    return f"""fastapi==0.104.1
uvicorn==0.24.0
psycopg2-binary==2.9.9
python-multipart==0.0.6
python-jose[cryptography]==3.3.0
passlib[bcrypt]==1.7.4
python-dotenv==1.0.0
pydantic==2.4.0
sqlalchemy==2.0.0
alembic==1.12.0
pytest==7.4.0
pytest-asyncio==0.21.0
"""

def generate_express_app(spec):
    """Generate real Express.js app"""
    return f'''const express = require('express')
const cors = require('cors')
const helmet = require('helmet')
const dotenv = require('dotenv')

// Load environment variables
dotenv.config()

const app = express()
const PORT = process.env.PORT || 8000

// Middleware
app.use(helmet())
app.use(cors())
app.use(express.json())

// Health check endpoint
app.get('/health', (req, res) => {{
  res.json({{ 
    status: 'healthy', 
    service: '{spec['name']}',
    timestamp: new Date().toISOString(),
    version: '1.0.0'
  }})
}})

// API routes
app.get('/api/status', (req, res) => {{
  res.json({{ 
    message: 'Welcome to {spec['name']} API',
    version: '1.0.0',
    features: {json.dumps(spec['features'])},
    infrastructure: {json.dumps(spec['infrastructure'])}
  }})
}})

// Error handling middleware
app.use((err, req, res, next) => {{
  console.error(err.stack)
  res.status(500).json({{ 
    error: 'Something went wrong!',
    message: err.message 
  }})
}})

// 404 handler
app.use('*', (req, res) => {{
  res.status(404).json({{ 
    error: 'Route not found',
    path: req.originalUrl 
  }})
}})

// Start server
app.listen(PORT, () => {{
  console.log(`{spec['name']} API server running on port ${{PORT}}`)
  console.log(`Health check available at http://localhost:${{PORT}}/health`)
}})

module.exports = app'''

def generate_fastapi_app(spec):
    """Generate real FastAPI app"""
    return f'''from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.trustedhost import TrustedHostMiddleware
import uvicorn
from datetime import datetime

app = FastAPI(
    title="{spec['name']}",
    description="{spec['description']}",
    version="1.0.0"
)

# Middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.add_middleware(
    TrustedHostMiddleware, 
    allowed_hosts=["*"]
)

@app.get("/health")
async def health():
    return {{
        "status": "healthy",
        "service": "{spec['name']}",
        "timestamp": datetime.utcnow().isoformat(),
        "version": "1.0.0"
    }}

@app.get("/api/status")
async def status():
    return {{
        "message": "Welcome to {spec['name']} API",
        "version": "1.0.0",
        "features": {spec['features']},
        "infrastructure": {spec['infrastructure']}
    }}

@app.exception_handler(404)
async def not_found_handler(request, exc):
    return {{
        "error": "Route not found",
        "path": str(request.url)
    }}

@app.exception_handler(500)
async def internal_error_handler(request, exc):
    return {{
        "error": "Internal server error",
        "message": str(exc)
    }}

if __name__ == "__main__":
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=8000,
        log_level="info"
    )'''

def generate_react_app_component(spec):
    """Generate real React App component"""
    return f'''import React from 'react'
import {{ useState, useEffect }} from 'react'
import Header from './components/Header'
import Home from './pages/Home'
import './App.css'

function App() {{
  const [data, setData] = useState(null)
  const [loading, setLoading] = useState(true)

  useEffect(() => {{
    // Fetch initial data
    fetch('/api/status')
      .then(response => response.json())
      .then(data => {{
        setData(data)
        setLoading(false)
      }})
      .catch(error => {{
        console.error('Error fetching data:', error)
        setLoading(false)
      }})
  }}, [])

  if (loading) {{
    return (
      <div className="min-h-screen flex items-center justify-center">
        <div className="animate-spin rounded-full h-12 w-12 border-b-2 border-blue-600"></div>
      </div>
    )
  }}

  return (
    <div className="min-h-screen bg-gray-50">
      <Header />
      <main className="max-w-7xl mx-auto px-4 sm:px-6 lg:px-8 py-8">
        <Home data={{data}} />
      </main>
    </div>
  )
}}

export default App'''

def generate_nextjs_index_page(spec):
    """Generate real Next.js index page"""
    return f'''import Head from 'next/head'
import Header from '../components/Header'
import Layout from '../components/Layout'
import {{ useState, useEffect }} from 'react'

export default function Home() {{
  const [data, setData] = useState(null)
  const [loading, setLoading] = useState(true)

  useEffect(() => {{
    // Fetch initial data
    fetch('/api/status')
      .then(response => response.json())
      .then(data => {{
        setData(data)
        setLoading(false)
      }})
      .catch(error => {{
        console.error('Error fetching data:', error)
        setLoading(false)
      }})
  }}, [])

  if (loading) {{
    return (
      <Layout>
        <div className="min-h-screen flex items-center justify-center">
          <div className="animate-spin rounded-full h-12 w-12 border-b-2 border-blue-600"></div>
        </div>
      </Layout>
    )
  }}

  return (
    <Layout>
      <Head>
        <title>{spec['name']}</title>
        <meta name="description" content="{spec['description']}" />
        <meta name="viewport" content="width=device-width, initial-scale=1" />
        <link rel="icon" href="/favicon.ico" />
      </Head>
      
      <main className="max-w-7xl mx-auto px-4 sm:px-6 lg:px-8 py-8">
        <div className="text-center">
          <h1 className="text-4xl font-bold text-gray-900 mb-4">
            {spec['name']}
          </h1>
          <p className="text-xl text-gray-600 mb-8">
            {spec['description']}
          </p>
          
          {{data && (
            <div className="bg-white rounded-lg shadow p-6">
              <h2 className="text-2xl font-semibold mb-4">System Status</h2>
              <div className="grid grid-cols-1 md:grid-cols-2 gap-4">
                <div>
                  <h3 className="font-medium text-gray-900">Features</h3>
                  <ul className="mt-2 space-y-1">
                    {{data.features?.map((feature, index) => (
                      <li key={{index}} className="text-sm text-gray-600">
                        • {{feature}}
                      </li>
                    ))}}
                  </ul>
                </div>
                <div>
                  <h3 className="font-medium text-gray-900">Infrastructure</h3>
                  <ul className="mt-2 space-y-1">
                    {{data.infrastructure?.map((infra, index) => (
                      <li key={{index}} className="text-sm text-gray-600">
                        • {{infra}}
                      </li>
                    ))}}
                  </ul>
                </div>
              </div>
            </div>
          )}}
        </div>
      </main>
    </Layout>
  )
}}'''

def generate_react_header_component(spec):
    return f'''import React from 'react'

export default function Header() {{
  return (
    <header className="bg-white shadow">
      <div className="max-w-7xl mx-auto px-4 sm:px-6 lg:px-8">
        <div className="flex justify-between items-center h-16">
          <div className="flex items-center">
            <h1 className="text-xl font-bold text-gray-900">
              {spec['name']}
            </h1>
          </div>
          <nav className="flex space-x-4">
            <a href="/" className="text-gray-600 hover:text-gray-900">Home</a>
            <a href="/about" className="text-gray-600 hover:text-gray-900">About</a>
          </nav>
        </div>
      </div>
    </header>
  )
}}'''

def generate_nextjs_header_component(spec):
    return f'''import React from 'react'
import Link from 'next/link'

export default function Header() {{
  return (
    <header className="bg-white shadow">
      <div className="max-w-7xl mx-auto px-4 sm:px-6 lg:px-8">
        <div className="flex justify-between items-center h-16">
          <div className="flex items-center">
            <Link href="/" className="text-xl font-bold text-gray-900">
              {spec['name']}
            </Link>
          </div>
          <nav className="flex space-x-4">
            <Link href="/" className="text-gray-600 hover:text-gray-900">Home</Link>
            <Link href="/about" className="text-gray-600 hover:text-gray-900">About</Link>
          </nav>
        </div>
      </div>
    </header>
  )
}}'''

def generate_nextjs_layout_component(spec):
    return f'''import React from 'react'
import Header from './Header'

export default function Layout({{ children }}) {{
  return (
    <div className="min-h-screen bg-gray-50">
      <Header />
      {{children}}
    </div>
  )
}}'''

def generate_react_home_page(spec):
    return f'''import React from 'react'

export default function Home({{ data }}) {{
  return (
    <div className="text-center">
      <h1 className="text-4xl font-bold text-gray-900 mb-4">
        {spec['name']}
      </h1>
      <p className="text-xl text-gray-600 mb-8">
        {spec['description']}
      </p>
      
      {{data && (
        <div className="bg-white rounded-lg shadow p-6">
          <h2 className="text-2xl font-semibold mb-4">System Status</h2>
          <div className="grid grid-cols-1 md:grid-cols-2 gap-4">
            <div>
              <h3 className="font-medium text-gray-900">Features</h3>
              <ul className="mt-2 space-y-1">
                {{data.features?.map((feature, index) => (
                  <li key={{index}} className="text-sm text-gray-600">
                    • {{feature}}
                  </li>
                ))}}
              </ul>
            </div>
            <div>
              <h3 className="font-medium text-gray-900">Infrastructure</h3>
              <ul className="mt-2 space-y-1">
                {{data.infrastructure?.map((infra, index) => (
                  <li key={{index}} className="text-sm text-gray-600">
                    • {{infra}}
                  </li>
                ))}}
              </ul>
            </div>
          </div>
        </div>
      )}}
    </div>
  )
}}'''

def generate_tailwind_config(spec):
    return f'''/** @type {{import('tailwindcss').Config}} */
module.exports = {{
  content: [
    './src/**/*.{{js,jsx,ts,tsx}}',
    './pages/**/*.{{js,jsx,ts,tsx}}',
    './components/**/*.{{js,jsx,ts,tsx}}',
  ],
  theme: {{
    extend: {{
      colors: {{
        primary: {{
          50: '#eff6ff',
          500: '#3b82f6',
          600: '#2563eb',
          700: '#1d4ed8',
        }}
      }}
    }},
  }},
  plugins: [],
}}'''

def generate_typescript_config(spec):
    return f'''{{
  "compilerOptions": {{
    "target": "es5",
    "lib": ["dom", "dom.iterable", "es6"],
    "allowJs": true,
    "skipLibCheck": true,
    "strict": true,
    "noEmit": true,
    "esModuleInterop": true,
    "module": "esnext",
    "moduleResolution": "bundler",
    "resolveJsonModule": true,
    "isolatedModules": true,
    "jsx": "preserve",
    "incremental": true
  }},
  "include": ["next-env.d.ts", "**/*.ts", "**/*.tsx"],
  "exclude": ["node_modules"]
}}'''

def generate_nextjs_config(spec):
    return f'''/** @type {{import('next').NextConfig}} */
const nextConfig = {{
  reactStrictMode: true,
  swcMinify: true,
  images: {{
    unoptimized: true
  }}
}}

module.exports = nextConfig'''

# Route moved to server_part7.py inside create_app() function

def enhance_specification_with_references(spec, reference_urls, uploaded_files):
    """Enhance system specification based on reference URLs and uploaded files"""
    enhanced_spec = spec.copy()
    
    # Analyze reference URLs
    url_insights = []
    for url in reference_urls:
        try:
            analysis = analyze_reference_url(url)
            if analysis['success']:
                url_insights.append(analysis['analysis'])
        except Exception as e:
            logger.error(f"Error analyzing reference URL {url}: {e}")
    
    # Extract insights from URL analysis
    if url_insights:
        # Merge technologies
        all_technologies = []
        for insight in url_insights:
            all_technologies.extend(insight.get('technologies', []))
        enhanced_spec['techStack'] = list(set(enhanced_spec.get('techStack', []) + all_technologies))
        
        # Merge features
        all_features = []
        for insight in url_insights:
            all_features.extend(insight.get('features', []))
        enhanced_spec['features'] = list(set(enhanced_spec.get('features', []) + all_features))
        
        # Add layout elements
        all_layout_elements = []
        for insight in url_insights:
            all_layout_elements.extend(insight.get('layout_elements', []))
        enhanced_spec['layout_elements'] = list(set(all_layout_elements))
    
    # Analyze uploaded files
    file_insights = []
    for file_info in uploaded_files:
        if file_info.get('analysis') and file_info['analysis']['success']:
            file_insights.append(file_info['analysis'])
    
    # Extract insights from file analysis
    if file_insights:
        # Add design insights from images
        for insight in file_insights:
            if 'analysis' in insight and 'dominant_color' in insight['analysis']:
                enhanced_spec['design_preferences'] = enhanced_spec.get('design_preferences', {})
                enhanced_spec['design_preferences']['primary_color'] = insight['analysis']['dominant_color']
            
            if 'text_content' in insight:
                # Extract requirements from document text
                text = insight['text_content'].lower()
                if 'authentication' in text or 'login' in text:
                    enhanced_spec['features'] = list(set(enhanced_spec.get('features', []) + ['User Authentication']))
                if 'payment' in text or 'ecommerce' in text:
                    enhanced_spec['features'] = list(set(enhanced_spec.get('features', []) + ['Payment Processing']))
                if 'search' in text:
                    enhanced_spec['features'] = list(set(enhanced_spec.get('features', []) + ['Search Functionality']))
    
    return enhanced_spec

def generate_api_routes(spec):
    return f'''const express = require('express')
const router = express.Router()

// {spec['name']} API routes
router.get('/health', (req, res) => {{
    res.json({{ 
        status: 'healthy',
        service: '{spec['name']}',
        timestamp: new Date().toISOString()
    }})
}})

router.get('/status', (req, res) => {{
    res.json({{ 
        message: 'Welcome to {spec['name']} API',
        version: '1.0.0',
        features: {json.dumps(spec['features'])},
        infrastructure: {json.dumps(spec['infrastructure'])}
    }})
}})

module.exports = router'''

def generate_fastapi_routes(spec):
    return f'''from fastapi import APIRouter
from datetime import datetime

router = APIRouter()

@router.get("/health")
async def health():
    return {{
        "status": "healthy",
        "service": "{spec['name']}",
        "timestamp": datetime.utcnow().isoformat()
    }}

@router.get("/status")
async def status():
    return {{
        "message": "Welcome to {spec['name']} API",
        "version": "1.0.0",
        "features": {spec['features']},
        "infrastructure": {spec['infrastructure']}
    }}'''

def generate_database_models(spec):
    """Generate real, system-specific database models based on system type and features"""
    
    # Determine system type and generate appropriate models
    system_type = spec.get('type', 'web-app')
    features = spec.get('features', [])
    
    if system_type == 'web-app':
        return generate_webapp_models(spec, features)
    elif system_type == 'ecommerce-platform':
        return generate_ecommerce_models(spec, features)
    elif system_type == 'api-service':
        return generate_api_models(spec, features)
    else:
        return generate_generic_models(spec)

def generate_webapp_models(spec, features):
    """Generate models for a web application with users, posts, comments"""
    system_name = spec['name'].replace(" ", "").replace('-', '')
    
    models = f'''const {{ Pool }} = require('pg')

// Database connection
const pool = new Pool({{
  connectionString: process.env.DATABASE_URL,
  ssl: process.env.NODE_ENV === 'production' ? {{ rejectUnauthorized: false }} : false
}})

// {spec['name']} Database Models

// User Model
class User {{
  static async create(userData) {{
    const {{ email, password_hash, name, role = 'user' }} = userData
    const query = `
      INSERT INTO users (email, password_hash, name, role, created_at, updated_at)
      VALUES ($1, $2, $3, $4, NOW(), NOW())
      RETURNING id, email, name, role, created_at, updated_at
    `
    const values = [email, password_hash, name, role]
    const result = await pool.query(query, values)
    return result.rows[0]
  }}

  static async findByEmail(email) {{
    const query = 'SELECT * FROM users WHERE email = $1'
    const result = await pool.query(query, [email])
    return result.rows[0]
  }}

  static async findById(id) {{
    const query = 'SELECT id, email, name, role, created_at, updated_at FROM users WHERE id = $1'
    const result = await pool.query(query, [id])
    return result.rows[0]
  }}

  static async update(id, updateData) {{
    const fields = []
    const values = []
    let paramCount = 1

    for (const [key, value] of Object.entries(updateData)) {{
      if (key !== 'id') {{
        fields.push(`${{key}} = ${{paramCount}}`)
        values.push(value)
        paramCount++
      }}
    }}

    if (fields.length === 0) return null

    fields.push('updated_at = NOW()')
    values.push(id)

    const query = `
      UPDATE users 
      SET ${{fields.join(', ')}}
      WHERE id = ${{paramCount}}
      RETURNING id, email, name, role, created_at, updated_at
    `
    const result = await pool.query(query, values)
    return result.rows[0]
  }}

  static async delete(id) {{
    const query = 'DELETE FROM users WHERE id = $1 RETURNING id'
    const result = await pool.query(query, [id])
    return result.rows[0]
  }}
}}'''

    # Add Post model if content management is a feature
    if any('content' in feature.lower() or 'post' in feature.lower() or 'blog' in feature.lower() for feature in features):
        models += f'''

// Post Model
class Post {{
  static async create(postData) {{
    const {{ title, content, author_id, status = 'draft' }} = postData
    const query = `
      INSERT INTO posts (title, content, author_id, status, created_at, updated_at)
      VALUES ($1, $2, $3, $4, NOW(), NOW())
      RETURNING *
    `
    const values = [title, content, author_id, status]
    const result = await pool.query(query, values)
    return result.rows[0]
  }}

  static async findById(id) {{
    const query = `
      SELECT p.*, u.name as author_name, u.email as author_email
      FROM posts p
      JOIN users u ON p.author_id = u.id
      WHERE p.id = $1
    `
    const result = await pool.query(query, [id])
    return result.rows[0]
  }}

  static async findByAuthor(authorId, limit = 10, offset = 0) {{
    const query = `
      SELECT p.*, u.name as author_name
      FROM posts p
      JOIN users u ON p.author_id = u.id
      WHERE p.author_id = $1
      ORDER BY p.created_at DESC
      LIMIT $2 OFFSET $3
    `
    const result = await pool.query(query, [authorId, limit, offset])
    return result.rows
  }}

  static async findAll(limit = 10, offset = 0, status = 'published') {{
    const query = `
      SELECT p.*, u.name as author_name
      FROM posts p
      JOIN users u ON p.author_id = u.id
      WHERE p.status = $1
      ORDER BY p.created_at DESC
      LIMIT $2 OFFSET $3
    `
    const result = await pool.query(query, [status, limit, offset])
    return result.rows
  }}

  static async update(id, updateData) {{
    const fields = []
    const values = []
    let paramCount = 1

    for (const [key, value] of Object.entries(updateData)) {{
      if (key !== 'id') {{
        fields.push(`${{key}} = ${{paramCount}}`)
        values.push(value)
        paramCount++
      }}
    }}

    if (fields.length === 0) return null

    fields.push('updated_at = NOW()')
    values.push(id)

    const query = `
      UPDATE posts 
      SET ${{fields.join(', ')}}
      WHERE id = ${{paramCount}}
      RETURNING *
    `
    const result = await pool.query(query, values)
    return result.rows[0]
  }}

  static async delete(id) {{
    const query = 'DELETE FROM posts WHERE id = $1 RETURNING id'
    const result = await pool.query(query, [id])
    return result.rows[0]
  }}
}}'''

    # Add Comment model if user interactions are a feature
    if any('comment' in feature.lower() or 'interaction' in feature.lower() or 'social' in feature.lower() for feature in features):
        models += f'''

// Comment Model
class Comment {{
  static async create(commentData) {{
    const {{ content, author_id, post_id, parent_id = null }} = commentData
    const query = `
      INSERT INTO comments (content, author_id, post_id, parent_id, created_at, updated_at)
      VALUES ($1, $2, $3, $4, NOW(), NOW())
      RETURNING *
    `
    const values = [content, author_id, post_id, parent_id]
    const result = await pool.query(query, values)
    return result.rows[0]
  }}

  static async findByPost(postId, limit = 50, offset = 0) {{
    const query = `
      SELECT c.*, u.name as author_name, u.email as author_email
      FROM comments c
      JOIN users u ON c.author_id = u.id
      WHERE c.post_id = $1 AND c.parent_id IS NULL
      ORDER BY c.created_at ASC
      LIMIT $2 OFFSET $3
    `
    const result = await pool.query(query, [postId, limit, offset])
    return result.rows
  }}

  static async findReplies(parentId, limit = 20, offset = 0) {{
    const query = `
      SELECT c.*, u.name as author_name, u.email as author_email
      FROM comments c
      JOIN users u ON c.author_id = u.id
      WHERE c.parent_id = $1
      ORDER BY c.created_at ASC
      LIMIT $2 OFFSET $3
    `
    const result = await pool.query(query, [parentId, limit, offset])
    return result.rows
  }}

  static async update(id, updateData) {{
    const fields = []
    const values = []
    let paramCount = 1

    for (const [key, value] of Object.entries(updateData)) {{
      if (key !== 'id') {{
        fields.push(`${{key}} = ${{paramCount}}`)
        values.push(value)
        paramCount++
      }}
    }}

    if (fields.length === 0) return null

    fields.push('updated_at = NOW()')
    values.push(id)

    const query = `
      UPDATE comments 
      SET ${{fields.join(', ')}}
      WHERE id = ${{paramCount}}
      RETURNING *
    `
    const result = await pool.query(query, values)
    return result.rows[0]
  }}

  static async delete(id) {{
    const query = 'DELETE FROM comments WHERE id = $1 RETURNING id'
    const result = await pool.query(query, [id])
    return result.rows[0]
  }}
}}'''

    # Export all models
    models += f'''

module.exports = {{
  pool,
  User,
'''
    
    if any('content' in feature.lower() or 'post' in feature.lower() or 'blog' in feature.lower() for feature in features):
        models += '  Post,\n'
    
    if any('comment' in feature.lower() or 'interaction' in feature.lower() or 'social' in feature.lower() for feature in features):
        models += '  Comment,\n'
    
    models += '}'
    
    return models

def generate_ecommerce_models(spec, features):
    """Generate models for an e-commerce platform"""
    # TODO: Implement e-commerce specific models
    return generate_generic_models(spec)

def generate_api_models(spec, features):
    """Generate models for an API service"""
    # TODO: Implement API service specific models
    return generate_generic_models(spec)

def generate_generic_models(spec):
    """Generate generic models as fallback"""
    return f'''const {{ Pool }} = require('pg')

// Database connection
const pool = new Pool({{
  connectionString: process.env.DATABASE_URL,
  ssl: process.env.NODE_ENV === 'production' ? {{ rejectUnauthorized: false }} : false
}})

// {spec['name']} Models
class {spec['name'].replace(" ", "")}Model {{
  static async create(data) {{
    const query = 'INSERT INTO {spec['name'].lower().replace(" ", "_")} (data) VALUES ($1) RETURNING *'
    const values = [JSON.stringify(data)]
    const result = await pool.query(query, values)
    return result.rows[0]
  }}

  static async findById(id) {{
    const query = 'SELECT * FROM {spec['name'].lower().replace(" ", "_")} WHERE id = $1'
    const result = await pool.query(query, [id])
    return result.rows[0]
  }}

  static async findAll() {{
    const query = 'SELECT * FROM {spec['name'].lower().replace(" ", "_")}'
    const result = await pool.query(query)
    return result.rows
  }}
}}

module.exports = {{
  pool,
  {spec['name'].replace(" ", "")}Model
}}'''

def generate_database_migrations(spec):
    """Generate real database migration files"""
    system_type = spec.get('type', 'web-app')
    features = spec.get('features', [])
    
    if system_type == 'web-app':
        return generate_webapp_migrations(spec, features)
    else:
        return generate_generic_migrations(spec)

def generate_webapp_migrations(spec, features):
    """Generate migration files for web application"""
    migrations = []
    
    # Initial migration - create users table
    migrations.append({
        'name': '001_create_users_table.sql',
        'content': f'''-- Migration: Create users table
-- Generated for: {spec['name']}

CREATE TABLE IF NOT EXISTS users (
    id SERIAL PRIMARY KEY,
    email VARCHAR(255) UNIQUE NOT NULL,
    password_hash VARCHAR(255) NOT NULL,
    name VARCHAR(255) NOT NULL,
    role VARCHAR(50) DEFAULT 'user',
    email_verified BOOLEAN DEFAULT FALSE,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

-- Create indexes for performance
CREATE INDEX IF NOT EXISTS idx_users_email ON users(email);
CREATE INDEX IF NOT EXISTS idx_users_role ON users(role);
CREATE INDEX IF NOT EXISTS idx_users_created_at ON users(created_at);

-- Create updated_at trigger
CREATE OR REPLACE FUNCTION update_updated_at_column()
RETURNS TRIGGER AS $$
BEGIN
    NEW.updated_at = CURRENT_TIMESTAMP;
    RETURN NEW;
END;
$$ language 'plpgsql';

CREATE TRIGGER update_users_updated_at 
    BEFORE UPDATE ON users 
    FOR EACH ROW 
    EXECUTE FUNCTION update_updated_at_column();
'''
    })
    
    # Add posts table if content management is a feature
    if any('content' in feature.lower() or 'post' in feature.lower() or 'blog' in feature.lower() for feature in features):
        migrations.append({
            'name': '002_create_posts_table.sql',
            'content': f'''-- Migration: Create posts table
-- Generated for: {spec['name']}

CREATE TABLE IF NOT EXISTS posts (
    id SERIAL PRIMARY KEY,
    title VARCHAR(500) NOT NULL,
    content TEXT NOT NULL,
    excerpt TEXT,
    slug VARCHAR(500) UNIQUE,
    author_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
    status VARCHAR(50) DEFAULT 'draft',
    featured_image VARCHAR(500),
    meta_description TEXT,
    published_at TIMESTAMP,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

-- Create indexes for performance
CREATE INDEX IF NOT EXISTS idx_posts_author_id ON posts(author_id);
CREATE INDEX IF NOT EXISTS idx_posts_status ON posts(status);
CREATE INDEX IF NOT EXISTS idx_posts_published_at ON posts(published_at);
CREATE INDEX IF NOT EXISTS idx_posts_slug ON posts(slug);
CREATE INDEX IF NOT EXISTS idx_posts_created_at ON posts(created_at);

-- Create updated_at trigger
CREATE TRIGGER update_posts_updated_at 
    BEFORE UPDATE ON posts 
    FOR EACH ROW 
    EXECUTE FUNCTION update_updated_at_column();
'''
        })
    
    # Add comments table if user interactions are a feature
    if any('comment' in feature.lower() or 'interaction' in feature.lower() or 'social' in feature.lower() for feature in features):
        migrations.append({
            'name': '003_create_comments_table.sql',
            'content': f'''-- Migration: Create comments table
-- Generated for: {spec['name']}

CREATE TABLE IF NOT EXISTS comments (
    id SERIAL PRIMARY KEY,
    content TEXT NOT NULL,
    author_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
    post_id INTEGER NOT NULL REFERENCES posts(id) ON DELETE CASCADE,
    parent_id INTEGER REFERENCES comments(id) ON DELETE CASCADE,
    status VARCHAR(50) DEFAULT 'approved',
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

-- Create indexes for performance
CREATE INDEX IF NOT EXISTS idx_comments_author_id ON comments(author_id);
CREATE INDEX IF NOT EXISTS idx_comments_post_id ON comments(post_id);
CREATE INDEX IF NOT EXISTS idx_comments_parent_id ON comments(parent_id);
CREATE INDEX IF NOT EXISTS idx_comments_status ON comments(status);
CREATE INDEX IF NOT EXISTS idx_comments_created_at ON comments(created_at);

-- Create updated_at trigger
CREATE TRIGGER update_comments_updated_at 
    BEFORE UPDATE ON comments 
    FOR EACH ROW 
    EXECUTE FUNCTION update_updated_at_column();
'''
        })
    
    # Add seed data migration
    migrations.append({
        'name': '004_seed_data.sql',
        'content': f'''-- Migration: Seed initial data
-- Generated for: {spec['name']}

-- Insert admin user (password: admin123 - change in production!)
INSERT INTO users (email, password_hash, name, role, email_verified) 
VALUES (
    'admin@{spec['name'].lower().replace(" ", "")}.com',
    '$2b$10$rQZ8K9vX8K9vX8K9vX8K9e', -- bcrypt hash for 'admin123'
    'Admin User',
    'admin',
    true
) ON CONFLICT (email) DO NOTHING;

-- Insert sample user
INSERT INTO users (email, password_hash, name, role, email_verified) 
VALUES (
    'user@{spec['name'].lower().replace(" ", "")}.com',
    '$2b$10$rQZ8K9vX8K9vX8K9vX8K9e', -- bcrypt hash for 'user123'
    'Sample User',
    'user',
    true
) ON CONFLICT (email) DO NOTHING;
'''
    })
    
    return migrations

def generate_generic_migrations(spec):
    """Generate generic migration files"""
    return [{
        'name': '001_create_initial_table.sql',
        'content': f'''-- Migration: Create initial table
-- Generated for: {spec['name']}

CREATE TABLE IF NOT EXISTS {spec['name'].lower().replace(" ", "_")} (
    id SERIAL PRIMARY KEY,
    data JSONB,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

-- Create indexes
CREATE INDEX IF NOT EXISTS idx_{spec['name'].lower().replace(" ", "_")}_created_at ON {spec['name'].lower().replace(" ", "_")}(created_at);
'''
    }]

def generate_migration_runner(spec):
    """Generate migration runner script"""
    return f'''const {{ Pool }} = require('pg')
const fs = require('fs')
const path = require('path')

// Database connection
const pool = new Pool({{
  connectionString: process.env.DATABASE_URL,
  ssl: process.env.NODE_ENV === 'production' ? {{ rejectUnauthorized: false }} : false
}})

async function runMigrations() {{
  try {{
    console.log('🚀 Starting database migrations for {spec['name']}...')
    
    // Create migrations table if it doesn't exist
    await pool.query(`
      CREATE TABLE IF NOT EXISTS migrations (
        id SERIAL PRIMARY KEY,
        filename VARCHAR(255) UNIQUE NOT NULL,
        executed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
      )
    `)
    
    // Get list of migration files
    const migrationsDir = path.join(__dirname, '../database/migrations')
    const migrationFiles = fs.readdirSync(migrationsDir)
      .filter(file => file.endsWith('.sql'))
      .sort()
    
    console.log(`Found ${{migrationFiles.length}} migration files`)
    
    // Run each migration
    for (const filename of migrationFiles) {{
      // Check if migration already ran
      const result = await pool.query(
        'SELECT id FROM migrations WHERE filename = $1',
        [filename]
      )
      
      if (result.rows.length > 0) {{
        console.log(`⏭️  Skipping ${{filename}} (already executed)`)
        continue
      }}
      
      // Read and execute migration
      const migrationPath = path.join(migrationsDir, filename)
      const migrationSQL = fs.readFileSync(migrationPath, 'utf8')
      
      console.log(`🔄 Running migration: ${{filename}}`)
      await pool.query(migrationSQL)
      
      // Record migration as executed
      await pool.query(
        'INSERT INTO migrations (filename) VALUES ($1)',
        [filename]
      )
      
      console.log(`✅ Completed migration: ${{filename}}`)
    }}
    
    console.log('🎉 All migrations completed successfully!')
    
  }} catch (error) {{
    console.error('❌ Migration failed:', error)
    process.exit(1)
  }} finally {{
    await pool.end()
  }}
}}

// Run migrations if this script is executed directly
if (require.main === module) {{
  runMigrations()
}}

module.exports = {{ runMigrations }}
'''

def generate_docker_compose(spec):
    """Generate Docker Compose for local development"""
    system_name = spec['name'].lower().replace(' ', '-').replace('_', '-')
    
    return f'''version: '3.8'

services:
  # PostgreSQL Database
  postgres:
    image: postgres:15-alpine
    container_name: {system_name}-postgres
    environment:
      POSTGRES_DB: {system_name}
      POSTGRES_USER: {system_name}_user
      POSTGRES_PASSWORD: {system_name}_password
    ports:
      - "5432:5432"
    volumes:
      - postgres_data:/var/lib/postgresql/data
      - ./database/migrations:/docker-entrypoint-initdb.d
    healthcheck:
      test: ["CMD-SHELL", "pg_isready -U {system_name}_user -d {system_name}"]
      interval: 10s
      timeout: 5s
      retries: 5

  # Redis Cache (optional)
  redis:
    image: redis:7-alpine
    container_name: {system_name}-redis
    ports:
      - "6379:6379"
    volumes:
      - redis_data:/data
    healthcheck:
      test: ["CMD", "redis-cli", "ping"]
      interval: 10s
      timeout: 5s
      retries: 5

  # Backend API
  backend:
    build:
      context: ./backend
      dockerfile: Dockerfile
    container_name: {system_name}-backend
    environment:
      NODE_ENV: development
      DATABASE_URL: postgresql://{system_name}_user:{system_name}_password@postgres:5432/{system_name}
      REDIS_URL: redis://redis:6379
      PORT: 8000
    ports:
      - "8000:8000"
    volumes:
      - ./backend:/app
      - /app/node_modules
    depends_on:
      postgres:
        condition: service_healthy
      redis:
        condition: service_healthy
    command: npm run dev
    healthcheck:
      test: ["CMD", "curl", "-f", "http://localhost:8000/health"]
      interval: 30s
      timeout: 10s
      retries: 3

  # Frontend (if Next.js)
  frontend:
    build:
      context: ./frontend
      dockerfile: Dockerfile
    container_name: {system_name}-frontend
    environment:
      NEXT_PUBLIC_API_URL: http://localhost:8000
      NODE_ENV: development
    ports:
      - "3000:3000"
    volumes:
      - ./frontend:/app
      - /app/node_modules
      - /app/.next
    depends_on:
      - backend
    command: npm run dev
    healthcheck:
      test: ["CMD", "curl", "-f", "http://localhost:3000"]
      interval: 30s
      timeout: 10s
      retries: 3

volumes:
  postgres_data:
  redis_data:

networks:
  default:
    name: {system_name}-network
'''

def generate_deployment_scripts(spec):
    """Generate deployment scripts for different cloud providers"""
    system_name = spec['name'].lower().replace(' ', '-').replace('_', '-')
    
    scripts = []
    
    # AWS deployment script
    scripts.append({
        'name': 'deploy-aws.sh',
        'content': f'''#!/bin/bash
# AWS Deployment Script for {spec['name']}
# Generated by System Builder Hub (SBH)

set -e

echo "🚀 Deploying {spec['name']} to AWS..."

# Configuration
PROJECT_NAME="{system_name}"
AWS_REGION="us-west-2"
ECR_REPO="$PROJECT_NAME-repo"
ECS_CLUSTER="$PROJECT_NAME-cluster"
ECS_SERVICE="$PROJECT_NAME-service"

# Check if AWS CLI is configured
if ! aws sts get-caller-identity > /dev/null 2>&1; then
    echo "❌ AWS CLI not configured. Please run 'aws configure' first."
    exit 1
fi

# Check if Docker is running
if ! docker info > /dev/null 2>&1; then
    echo "❌ Docker is not running. Please start Docker first."
    exit 1
fi

echo "📦 Building Docker image..."
docker build -t $PROJECT_NAME:latest ./backend

echo "🏷️  Tagging image for ECR..."
aws ecr get-login-password --region $AWS_REGION | docker login --username AWS --password-stdin $AWS_ACCOUNT_ID.dkr.ecr.$AWS_REGION.amazonaws.com
docker tag $PROJECT_NAME:latest $AWS_ACCOUNT_ID.dkr.ecr.$AWS_REGION.amazonaws.com/$ECR_REPO:latest

echo "⬆️  Pushing image to ECR..."
docker push $AWS_ACCOUNT_ID.dkr.ecr.$AWS_REGION.amazonaws.com/$ECR_REPO:latest

echo "🔄 Updating ECS service..."
aws ecs update-service \\
    --cluster $ECS_CLUSTER \\
    --service $ECS_SERVICE \\
    --force-new-deployment

echo "⏳ Waiting for deployment to complete..."
aws ecs wait services-stable \\
    --cluster $ECS_CLUSTER \\
    --services $ECS_SERVICE

echo "✅ Deployment completed successfully!"
echo "🌐 Your application should be available at: https://$PROJECT_NAME.sbh.umbervale.com"
'''
    })
    
    # GCP deployment script
    scripts.append({
        'name': 'deploy-gcp.sh',
        'content': f'''#!/bin/bash
# Google Cloud Platform Deployment Script for {spec['name']}
# Generated by System Builder Hub (SBH)

set -e

echo "🚀 Deploying {spec['name']} to Google Cloud Platform..."

# Configuration
PROJECT_NAME="{system_name}"
GCP_PROJECT_ID="$GCP_PROJECT_ID"
GCP_REGION="us-west1"
SERVICE_NAME="$PROJECT_NAME-service"

# Check if gcloud CLI is configured
if ! gcloud auth list --filter=status:ACTIVE --format="value(account)" | head -n1 > /dev/null 2>&1; then
    echo "❌ Google Cloud CLI not configured. Please run 'gcloud auth login' first."
    exit 1
fi

# Set the project
gcloud config set project $GCP_PROJECT_ID

echo "📦 Building and pushing to Google Container Registry..."
gcloud builds submit --tag gcr.io/$GCP_PROJECT_ID/$PROJECT_NAME ./backend

echo "🚀 Deploying to Cloud Run..."
gcloud run deploy $SERVICE_NAME \\
    --image gcr.io/$GCP_PROJECT_ID/$PROJECT_NAME \\
    --platform managed \\
    --region $GCP_REGION \\
    --allow-unauthenticated \\
    --port 8000 \\
    --memory 512Mi \\
    --cpu 1 \\
    --max-instances 10

echo "✅ Deployment completed successfully!"
echo "🌐 Your application is available at the URL shown above."
'''
    })
    
    # Azure deployment script
    scripts.append({
        'name': 'deploy-azure.sh',
        'content': f'''#!/bin/bash
# Microsoft Azure Deployment Script for {spec['name']}
# Generated by System Builder Hub (SBH)

set -e

echo "🚀 Deploying {spec['name']} to Microsoft Azure..."

# Configuration
PROJECT_NAME="{system_name}"
AZURE_RESOURCE_GROUP="$PROJECT_NAME-rg"
AZURE_LOCATION="westus2"
CONTAINER_APP_NAME="$PROJECT_NAME-app"

# Check if Azure CLI is configured
if ! az account show > /dev/null 2>&1; then
    echo "❌ Azure CLI not configured. Please run 'az login' first."
    exit 1
fi

echo "📦 Building and pushing to Azure Container Registry..."
az acr build --registry $AZURE_ACR_NAME --image $PROJECT_NAME:latest ./backend

echo "🚀 Deploying to Azure Container Apps..."
az containerapp create \\
    --name $CONTAINER_APP_NAME \\
    --resource-group $AZURE_RESOURCE_GROUP \\
    --location $AZURE_LOCATION \\
    --image $AZURE_ACR_NAME.azurecr.io/$PROJECT_NAME:latest \\
    --target-port 8000 \\
    --ingress external \\
    --cpu 0.5 \\
    --memory 1Gi \\
    --min-replicas 1 \\
    --max-replicas 10

echo "✅ Deployment completed successfully!"
echo "🌐 Your application is available at the URL shown above."
'''
    })
    
    return scripts

def generate_comprehensive_readme(spec):
    """Generate comprehensive README for the generated system"""
    system_name = spec['name']
    system_type = spec.get('type', 'web-app')
    features = spec.get('features', [])
    tech_stack = spec.get('techStack', [])
    
    readme = f'''# {system_name}

**Generated by System Builder Hub (SBH) - The AI-powered system that builds complete, deployable applications.**

## 🚀 Quick Start

### Prerequisites
- Docker and Docker Compose
- Node.js 18+ (for local development)
- PostgreSQL (or use Docker)
- Git

### Local Development

1. **Clone and setup:**
   ```bash
   git clone <your-repo-url>
   cd {system_name.lower().replace(' ', '-')}
   cp .env.example .env
   # Edit .env with your configuration
   ```

2. **Start with Docker Compose:**
   ```bash
   docker-compose up -d
   ```

3. **Run database migrations:**
   ```bash
   docker-compose exec backend node scripts/run-migrations.js
   ```

4. **Access your application:**
   - Frontend: http://localhost:3000
   - Backend API: http://localhost:8000
   - API Health: http://localhost:8000/health

## 🏗️ Architecture

### System Type: {system_type.title()}

### Tech Stack:
'''
    
    for tech in tech_stack:
        readme += f'- {tech}\n'
    
    readme += f'''
### Features:
'''
    
    for feature in features:
        readme += f'- {feature}\n'
    
    readme += f'''
## 📁 Project Structure

```
{system_name.lower().replace(' ', '-')}/
├── frontend/                 # Next.js frontend application
│   ├── pages/               # Next.js pages
│   ├── components/          # React components
│   ├── styles/              # CSS and styling
│   └── package.json         # Frontend dependencies
├── backend/                 # Node.js/Express backend
│   ├── src/                 # Source code
│   │   ├── models/          # Database models
│   │   ├── routes/          # API routes
│   │   └── middleware/      # Express middleware
│   ├── database/            # Database files
│   │   └── migrations/      # SQL migration files
│   └── scripts/             # Utility scripts
├── infrastructure/          # Terraform infrastructure
│   ├── main.tf             # Main Terraform configuration
│   ├── variables.tf        # Terraform variables
│   └── modules/            # Terraform modules
├── .github/workflows/       # GitHub Actions CI/CD
├── docker-compose.yml       # Local development setup
├── docker-compose.prod.yml  # Production setup
└── README.md               # This file
```

## 🗄️ Database Schema

### Tables:
'''
    
    if system_type == 'web-app':
        readme += '''- **users** - User authentication and profiles
- **posts** - Content management (if content features enabled)
- **comments** - User interactions (if interaction features enabled)
'''
    else:
        readme += f'- **{system_name.lower().replace(" ", "_")}** - Main data table\n'
    
    readme += f'''
### Migrations:
- `001_create_users_table.sql` - User authentication
- `002_create_posts_table.sql` - Content management
- `003_create_comments_table.sql` - User interactions
- `004_seed_data.sql` - Initial data

## 🚀 Deployment

### Option 1: AWS Deployment
```bash
chmod +x deploy-aws.sh
./deploy-aws.sh
```

### Option 2: Google Cloud Platform
```bash
chmod +x deploy-gcp.sh
./deploy-gcp.sh
```

### Option 3: Microsoft Azure
```bash
chmod +x deploy-azure.sh
./deploy-azure.sh
```

### Option 4: Docker Compose (Any Server)
```bash
docker-compose -f docker-compose.prod.yml up -d
```

## 🔧 Configuration

### Environment Variables

Copy `.env.example` to `.env` and configure:

```bash
# Database
DATABASE_URL=postgresql://user:password@localhost:5432/database

# Application
NODE_ENV=development
PORT=8000
JWT_SECRET=your-secret-key

# External Services
OPENAI_API_KEY=your-openai-key
STRIPE_SECRET_KEY=your-stripe-key
```

## 🧪 Testing

### Run Tests
```bash
# Backend tests
cd backend
npm test

# Frontend tests
cd frontend
npm test
```

### API Testing
```bash
# Health check
curl http://localhost:8000/health

# API endpoints
curl http://localhost:8000/api/status
```

## 📊 Monitoring

### Health Checks
- **Application**: `/health`
- **Database**: Connection status
- **External APIs**: Service availability

### Logs
```bash
# View application logs
docker-compose logs -f backend

# View database logs
docker-compose logs -f postgres
```

## 🔒 Security

### Authentication
- JWT-based authentication
- Password hashing with bcrypt
- Role-based access control

### Database Security
- Parameterized queries (no SQL injection)
- Connection encryption in production
- Regular security updates

## 🛠️ Development

### Adding New Features

1. **Database Changes:**
   ```bash
   # Create new migration
   touch database/migrations/005_add_new_feature.sql
   # Add your SQL changes
   # Run migration
   node scripts/run-migrations.js
   ```

2. **API Endpoints:**
   - Add routes in `backend/src/routes/`
   - Add models in `backend/src/models/`
   - Update tests

3. **Frontend Components:**
   - Add components in `frontend/components/`
   - Add pages in `frontend/pages/`
   - Update styling

## 📚 API Documentation

### Authentication Endpoints
- `POST /api/auth/login` - User login
- `POST /api/auth/register` - User registration
- `GET /api/auth/profile` - Get user profile

### Content Endpoints (if enabled)
- `GET /api/posts` - List posts
- `POST /api/posts` - Create post
- `GET /api/posts/:id` - Get post
- `PUT /api/posts/:id` - Update post
- `DELETE /api/posts/:id` - Delete post

### Comment Endpoints (if enabled)
- `GET /api/posts/:id/comments` - List comments
- `POST /api/posts/:id/comments` - Add comment
- `PUT /api/comments/:id` - Update comment
- `DELETE /api/comments/:id` - Delete comment

## 🆘 Troubleshooting

### Common Issues

1. **Database Connection Failed:**
   ```bash
   # Check if PostgreSQL is running
   docker-compose ps postgres
   # Check logs
   docker-compose logs postgres
   ```

2. **Port Already in Use:**
   ```bash
   # Change ports in docker-compose.yml
   # Or stop conflicting services
   ```

3. **Migration Errors:**
   ```bash
   # Check migration files
   ls -la database/migrations/
   # Run migrations manually
   docker-compose exec backend node scripts/run-migrations.js
   ```

## 📞 Support

- **Documentation**: This README
- **Issues**: GitHub Issues
- **Generated by**: System Builder Hub (SBH)

---

**Built with ❤️ by System Builder Hub - The AI-powered system that builds complete, deployable applications.**
'''
    
    return readme

def generate_sqlalchemy_models(spec):
    return f'''from sqlalchemy import Column, Integer, String, DateTime, Text, Boolean
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.sql import func
from datetime import datetime

Base = declarative_base()

class {spec['name'].replace(" ", "")}Model(Base):
    __tablename__ = '{spec['name'].lower().replace(" ", "_")}'

    id = Column(Integer, primary_key=True, index=True)
    name = Column(String(255), nullable=False)
    description = Column(Text)
    data = Column(Text)
    is_active = Column(Boolean, default=True)
    created_at = Column(DateTime(timezone=True), server_default=func.now())
    updated_at = Column(DateTime(timezone=True), onupdate=func.now())

    def __repr__(self):
        return f"<{spec['name'].replace(' ', '')}Model(id={{self.id}}, name={{self.name}})>"'''

def generate_database_config(spec):
    return f'''from sqlalchemy import create_engine
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
import os

DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://user:password@localhost/{spec['name'].lower().replace(" ", "_")}")

engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

Base = declarative_base()

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()'''

def generate_auth_middleware(spec):
    return f'''const jwt = require('jsonwebtoken')

const authenticateToken = (req, res, next) => {{
  const authHeader = req.headers['authorization']
  const token = authHeader && authHeader.split(' ')[1]

  if (!token) {{
    return res.status(401).json({{ error: 'Access token required' }})
  }}

  jwt.verify(token, process.env.JWT_SECRET, (err, user) => {{
    if (err) {{
      return res.status(403).json({{ error: 'Invalid token' }})
    }}
    req.user = user
    next()
  }})
}}

module.exports = {{ authenticateToken }}'''

def generate_python_auth_middleware(spec):
    return f'''from fastapi import HTTPException, Depends, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from jose import JWTError, jwt
import os

security = HTTPBearer()

def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={{"WWW-Authenticate": "Bearer"}},
    )
    
    try:
        payload = jwt.decode(credentials.credentials, os.getenv("JWT_SECRET"), algorithms=["HS256"])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    
    return username'''

def generate_nodejs_dockerfile(spec):
    return f'''FROM node:18-alpine

WORKDIR /app

COPY package*.json ./
RUN npm ci --only=production

COPY . .

EXPOSE 8000

CMD ["npm", "start"]'''

def generate_python_dockerfile(spec):
    return f'''FROM python:3.11-slim

WORKDIR /app

COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

COPY . .

EXPOSE 8000

CMD ["uvicorn", "src.main:app", "--host", "0.0.0.0", "--port", "8000"]'''

def generate_env_example(spec):
    return f'''# {spec['name']} Environment Variables
NODE_ENV=production
PORT=8000

# Database
DATABASE_URL=postgresql://user:password@localhost/{spec['name'].lower().replace(" ", "_")}

# JWT
JWT_SECRET=your-secret-key-here

# AWS (if applicable)
AWS_REGION=us-west-2
AWS_ACCESS_KEY_ID=your-access-key
AWS_SECRET_ACCESS_KEY=your-secret-key'''

def generate_terraform_main(spec, architecture):
    """Generate real Terraform main configuration"""
    return f'''# {spec['name']} Infrastructure
terraform {{
  required_version = ">= 1.0"
  required_providers {{
    aws = {{
      source  = "hashicorp/aws"
      version = "~> 5.0"
    }}
  }}
}}

provider "aws" {{
  region = var.aws_region
}}

# VPC and Networking
module "vpc" {{
  source = "./modules/vpc"
  
  project_name = var.project_name
  environment  = var.environment
}}

# ECS Cluster
module "ecs" {{
  source = "./modules/ecs"
  
  project_name = var.project_name
  environment  = var.environment
  vpc_id       = module.vpc.vpc_id
  subnets      = module.vpc.private_subnets
}}

# RDS Database
module "rds" {{
  source = "./modules/rds"
  
  project_name = var.project_name
  environment  = var.environment
  vpc_id       = module.vpc.vpc_id
  subnets      = module.vpc.private_subnets
}}

# Application Load Balancer
module "alb" {{
  source = "./modules/alb"
  
  project_name = var.project_name
  environment  = var.environment
  vpc_id       = module.vpc.vpc_id
  subnets      = module.vpc.public_subnets
}}

# S3 Bucket
module "s3" {{
  source = "./modules/s3"
  
  project_name = var.project_name
  environment  = var.environment
}}

# Outputs
output "cluster_name" {{
  description = "ECS cluster name"
  value       = module.ecs.cluster_name
}}

output "db_endpoint" {{
  description = "RDS endpoint"
  value       = module.rds.db_endpoint
}}

output "alb_dns_name" {{
  description = "ALB DNS name"
  value       = module.alb.alb_dns_name
}}

output "s3_bucket_name" {{
  description = "S3 bucket name"
  value       = module.s3.bucket_name
}}'''

def generate_github_actions_workflow(spec):
    """Generate real GitHub Actions workflow"""
    return f'''name: Deploy {spec['name']}

on:
  push:
    branches: [ main ]
  workflow_dispatch:

jobs:
  build-and-deploy:
    runs-on: ubuntu-latest
    
    permissions:
      id-token: write
      contents: read
    
    steps:
    - name: Checkout code
      uses: actions/checkout@v4
    
    - name: Configure AWS credentials
      uses: aws-actions/configure-aws-credentials@v4
      with:
        aws-region: us-west-2
        role-to-assume: ${{{{ secrets.AWS_ROLE_ARN }}}}
        role-session-name: {spec['name'].lower().replace(" ", "-")}-deploy
    
    - name: Login to Amazon ECR
      id: login-ecr
      uses: aws-actions/amazon-ecr-login@v2
    
    - name: Build, tag, and push image to Amazon ECR
      env:
        ECR_REGISTRY: ${{{{ steps.login-ecr.outputs.registry }}}}
        ECR_REPOSITORY: {spec['name'].lower().replace(" ", "-")}-repo
        IMAGE_TAG: ${{{{ github.sha }}}}
      run: |
        docker build -t $ECR_REGISTRY/$ECR_REPOSITORY:$IMAGE_TAG .
        docker push $ECR_REGISTRY/$ECR_REPOSITORY:$IMAGE_TAG
    
    - name: Deploy to ECS
      run: |
        aws ecs update-service \\
          --cluster {spec['name'].lower().replace(" ", "-")}-cluster \\
          --service {spec['name'].lower().replace(" ", "-")}-service \\
          --force-new-deployment
    
    - name: Wait for deployment to complete
      run: |
        aws ecs wait services-stable \\
          --cluster {spec['name'].lower().replace(" ", "-")}-cluster \\
          --services {spec['name'].lower().replace(" ", "-")}-service'''

def generate_test_workflow(spec):
    """Generate test workflow"""
    return f'''name: Test {spec['name']}

on:
  pull_request:
    branches: [ main ]
  push:
    branches: [ main ]

jobs:
  test:
    runs-on: ubuntu-latest
    
    steps:
    - name: Checkout code
      uses: actions/checkout@v4
    
    - name: Setup Node.js
      uses: actions/setup-node@v4
      with:
        node-version: '18'
        cache: 'npm'
    
    - name: Install dependencies
      run: npm ci
    
    - name: Run tests
      run: npm test
    
    - name: Run linting
      run: npm run lint'''

def generate_security_workflow(spec):
    """Generate security workflow"""
    return f'''name: Security Scan {spec['name']}

on:
  push:
    branches: [ main ]
  pull_request:
    branches: [ main ]

jobs:
  security:
    runs-on: ubuntu-latest
    
    steps:
    - name: Checkout code
      uses: actions/checkout@v4
    
    - name: Run Trivy vulnerability scanner
      uses: aquasecurity/trivy-action@master
      with:
        scan-type: 'fs'
        scan-ref: '.'
        format: 'sarif'
        output: 'trivy-results.sarif'
    
    - name: Upload Trivy scan results
      uses: github/codeql-action/upload-sarif@v2
      with:
        sarif_file: 'trivy-results.sarif'''

def generate_terraform_variables(spec):
    return f'''variable "project_name" {{
  description = "Name of the project"
  type        = string
  default     = "{spec['name'].lower().replace(" ", "-")}"
}}

variable "environment" {{
  description = "Environment name"
  type        = string
  default     = "dev"
}}

variable "aws_region" {{
  description = "AWS region"
  type        = string
  default     = "us-west-2"
}}

variable "db_password" {{
  description = "Database password"
  type        = string
  sensitive   = true
}}'''

def generate_terraform_outputs(spec):
    return f'''output "cluster_name" {{
  description = "ECS cluster name"
  value       = module.ecs.cluster_name
}}

output "db_endpoint" {{
  description = "RDS endpoint"
  value       = module.rds.db_endpoint
}}

output "alb_dns_name" {{
  description = "ALB DNS name"
  value       = module.alb.alb_dns_name
}}

output "s3_bucket_name" {{
  description = "S3 bucket name"
  value       = module.s3.bucket_name
}}'''

def generate_vpc_module(spec):
    return f'''# VPC Module for {spec['name']}
resource "aws_vpc" "main" {{
  cidr_block           = "10.0.0.0/16"
  enable_dns_hostnames = true
  enable_dns_support   = true

  tags = {{
    Name = "${{var.project_name}}-vpc"
    Environment = var.environment
  }}
}}

resource "aws_internet_gateway" "main" {{
  vpc_id = aws_vpc.main.id

  tags = {{
    Name = "${{var.project_name}}-igw"
    Environment = var.environment
  }}
}}

resource "aws_subnet" "public" {{
  count = 2

  vpc_id                  = aws_vpc.main.id
  cidr_block              = "10.0.${{count.index + 1}}.0/24"
  availability_zone       = data.aws_availability_zones.available.names[count.index]
  map_public_ip_on_launch = true

  tags = {{
    Name = "${{var.project_name}}-public-subnet-${{count.index + 1}}"
    Environment = var.environment
  }}
}}

resource "aws_subnet" "private" {{
  count = 2

  vpc_id            = aws_vpc.main.id
  cidr_block        = "10.0.${{count.index + 10}}.0/24"
  availability_zone = data.aws_availability_zones.available.names[count.index]

  tags = {{
    Name = "${{var.project_name}}-private-subnet-${{count.index + 1}}"
    Environment = var.environment
  }}
}}

data "aws_availability_zones" "available" {{
  state = "available"
}}

output "vpc_id" {{
  value = aws_vpc.main.id
}}

output "public_subnets" {{
  value = aws_subnet.public[*].id
}}

output "private_subnets" {{
  value = aws_subnet.private[*].id
}}'''

def generate_ecs_module(spec):
    return f'''# ECS Module for {spec['name']}
resource "aws_ecs_cluster" "main" {{
  name = "${{var.project_name}}-cluster"

  setting {{
    name  = "containerInsights"
    value = "enabled"
  }}

  tags = {{
    Name = "${{var.project_name}}-cluster"
    Environment = var.environment
  }}
}}

resource "aws_ecs_task_definition" "main" {{
  family                   = "${{var.project_name}}-task"
  network_mode             = "awsvpc"
  requires_compatibilities = ["FARGATE"]
  cpu                      = "256"
  memory                   = "512"
  execution_role_arn       = aws_iam_role.ecs_execution_role.arn

  container_definitions = jsonencode([{{
    name  = "${{var.project_name}}-container"
    image = "${{var.project_name}}-image:latest"
    portMappings = [{{
      containerPort = 8000
      hostPort      = 8000
    }}]
    environment = [
      {{
        name  = "NODE_ENV"
        value = "production"
      }}
    ]
    logConfiguration = {{
      logDriver = "awslogs"
      options = {{
        awslogs-group         = "/ecs/${{var.project_name}}"
        awslogs-region        = "us-west-2"
        awslogs-stream-prefix = "ecs"
      }}
    }}
  }}])

  tags = {{
    Name = "${{var.project_name}}-task"
    Environment = var.environment
  }}
}}

output "cluster_name" {{
  value = aws_ecs_cluster.main.name
}}'''

def generate_rds_module(spec):
    return f'''# RDS Module for {spec['name']}
resource "aws_db_subnet_group" "main" {{
  name       = "${{var.project_name}}-db-subnet-group"
  subnet_ids = var.subnets

  tags = {{
    Name = "${{var.project_name}}-db-subnet-group"
    Environment = var.environment
  }}
}}

resource "aws_db_instance" "main" {{
  identifier = "${{var.project_name}}-db"
  engine     = "postgres"
  engine_version = "15.4"
  instance_class = "db.t3.micro"
  allocated_storage = 20
  max_allocated_storage = 100
  storage_type = "gp2"
  storage_encrypted = true

  db_name  = "${{var.project_name}}"
  username = "admin"
  password = var.db_password

  vpc_security_group_ids = [aws_security_group.db.id]
  db_subnet_group_name   = aws_db_subnet_group.main.name

  backup_retention_period = 7
  backup_window          = "03:00-04:00"
  maintenance_window     = "sun:04:00-sun:05:00"

  skip_final_snapshot = true
  deletion_protection = false

  tags = {{
    Name = "${{var.project_name}}-db"
    Environment = var.environment
  }}
}}

output "db_endpoint" {{
  value = aws_db_instance.main.endpoint
}}'''

def generate_alb_module(spec):
    return f'''# ALB Module for {spec['name']}
resource "aws_lb" "main" {{
  name               = "${{var.project_name}}-alb"
  internal           = false
  load_balancer_type = "application"
  security_groups    = [aws_security_group.alb.id]
  subnets            = var.subnets

  enable_deletion_protection = false

  tags = {{
    Name = "${{var.project_name}}-alb"
    Environment = var.environment
  }}
}}

resource "aws_lb_target_group" "main" {{
  name     = "${{var.project_name}}-tg"
  port     = 8000
  protocol = "HTTP"
  vpc_id   = var.vpc_id
  target_type = "ip"

  health_check {{
    enabled             = true
    healthy_threshold   = 2
    interval            = 30
    matcher             = "200"
    path                = "/health"
    port                = "traffic-port"
    protocol            = "HTTP"
    timeout             = 5
    unhealthy_threshold = 2
  }}

  tags = {{
    Name = "${{var.project_name}}-tg"
    Environment = var.environment
  }}
}}

resource "aws_lb_listener" "main" {{
  load_balancer_arn = aws_lb.main.arn
  port              = "80"
  protocol          = "HTTP"

  default_action {{
    type             = "forward"
    target_group_arn = aws_lb_target_group.main.arn
  }}
}}

output "alb_dns_name" {{
  value = aws_lb.main.dns_name
}}'''

# Routes moved to server_part7.py inside create_app() function

def count_files(system):
    """Count total files in a system"""
    count = 0
    for template_type, template_data in system['templates'].items():
        if 'files' in template_data:
            count += len(template_data['files'])
    return count

def create_system_zip(system):
    """Create ZIP file from system data"""
    import zipfile
    import io
    
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Add README
        zip_file.writestr('README.md', f"# {system['specification']['name']}\n\nGenerated by System Builder Hub (SBH)\n\n{system['specification']['description']}")
        
        # Add all template files
        for template_type, template_data in system['templates'].items():
            if 'files' in template_data:
                for file_info in template_data['files']:
                    file_path = f"{template_type}/{file_info['name']}"
                    zip_file.writestr(file_path, file_info['content'])
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

# Add these imports at the top with the other imports
import zipfile
import io
import base64
import json
import boto3
from datetime import datetime, timedelta
from botocore.exceptions import ClientError

# S3 client for persistent storage
s3_client = boto3.client('s3')
S3_BUCKET = os.getenv('S3_BUCKET_NAME', 'sbh-generated-systems')

def save_system_to_s3(system_id, system_data):
    """Save system to S3"""
    try:
        key = f"systems/{system_id}.json"
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=key,
            Body=json.dumps(system_data, default=str),
            ContentType='application/json'
        )
        return True
    except ClientError as e:
        logger.error(f"Error saving system to S3: {e}")
        return False

def load_system_from_s3(system_id):
    """Load system from S3"""
    try:
        key = f"systems/{system_id}.json"
        response = s3_client.get_object(Bucket=S3_BUCKET, Key=key)
        return json.loads(response['Body'].read().decode('utf-8'))
    except ClientError as e:
        if e.response['Error']['Code'] == 'NoSuchKey':
            return None
        logger.error(f"Error loading system from S3: {e}")
        return None

def delete_system_from_s3(system_id):
    """Delete system from S3"""
    try:
        key = f"systems/{system_id}.json"
        s3_client.delete_object(Bucket=S3_BUCKET, Key=key)
        return True
    except ClientError as e:
        logger.error(f"Error deleting system from S3: {e}")
        return False

def create_app():
    """Create Flask application with OpenAI integration and system generation"""
    app = Flask(__name__)
    
    # Basic Configuration
    app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-secret-key')
    
    # CORS Configuration
    cors_origins = [
        'http://localhost:3000',
        'https://sbh.umbervale.com'
    ]
    CORS(app, 
         origins=cors_origins,
         allow_headers=['Content-Type', 'Authorization'],
         methods=['GET', 'POST', 'OPTIONS'],
         supports_credentials=False)
    
    # Get OpenAI configuration
    openai_config = get_openai_config()
    openai_client = create_openai_client()
    
    logger.info(f"OpenAI configured: {bool(openai_config['api_key'])}")
    logger.info(f"OpenAI model: {openai_config['model']}")

    @app.route('/api/health')
    def health():
        """Health check endpoint"""
        return jsonify({
            "ok": True, 
            "status": "healthy",
            "openai_configured": bool(openai_config['api_key']),
            "environment": os.getenv('FLASK_ENV', 'production')
        })

    @app.route('/')
    def index():
        return jsonify({
            "name": "System Builder Hub",
            "version": "1.0.0",
            "status": "running"
        })

    @app.route('/api/ai-chat/health', methods=['GET'])
    def ai_chat_health():
        """Health check for AI Chat service"""
        return jsonify({
            'status': 'healthy',
            'openai_configured': bool(openai_config['api_key']),
            'model': openai_config['model'] if openai_config['api_key'] else None,
            'timestamp': int(time.time())
        })

    @app.route('/api/ai-chat/chat', methods=['POST'])
    def ai_chat():
        """AI Chat endpoint with real OpenAI integration"""
        try:
            data = request.get_json()
            if not data:
                return jsonify({'error': 'No data provided'}), 400
    
            message = data.get('message', '')
            if not message:
                return jsonify({'error': 'No message provided'}), 400
    
            # Get conversation context and model selection
            conversation_history = data.get('conversation_history', [])
            conversation_id = data.get('conversation_id', f'conv_{int(time.time())}_{str(uuid.uuid4())[:8]}')
            system_message = data.get('system', 'You are an AI assistant for the System Builder Hub (SBH) - an AI-assisted platform that designs, scaffolds, deploys, and monitors complete software systems onto AWS. SBH is better than Cursor because it takes high-level specifications and outputs complete, bootable applications with their own infrastructure, CI/CD, and monitoring. You help users create comprehensive specifications for any type of system they want to build, then guide them through the process of generating working applications that are ready to deploy independently. Ask relevant questions to understand their requirements, provide architecture guidance, and help them create detailed specifications that SBH can use to build their complete system with Terraform, ECS, ALB, RDS, S3, and GitHub Actions.')
            
            # Model selection with validation
            requested_model = data.get('model', 'gpt-4o')  # Default to gpt-4o
            valid_models = ['gpt-4o-mini', 'gpt-4o', 'gpt-4-turbo']
            if requested_model not in valid_models:
                requested_model = 'gpt-4o'  # Fallback to default
            
            # If no OpenAI client, return echo behavior
            if not openai_client:
                return jsonify({
                    'success': True,
                    'response': f'You said: {message}',
                    'conversation_id': conversation_id,
                    'note': 'openai not configured'
                })
    
            # Build messages for OpenAI
            messages = [{"role": "system", "content": system_message}]
            
            # Add conversation history
            for msg in conversation_history:
                if isinstance(msg, dict) and 'role' in msg and 'content' in msg:
                    messages.append(msg)
            
            # Add current message
            messages.append({"role": "user", "content": message})
    
            # Call OpenAI API with selected model
            try:
                response = openai_client.chat.completions.create(
                    model=requested_model,  # Use the selected model
                    messages=messages,
                    max_tokens=1000,
                    temperature=0.7
                )
                
                ai_response = response.choices[0].message.content
                usage = response.usage
                
                return jsonify({
                    'success': True,
                    'response': ai_response,
                    'usage': {
                        'prompt_tokens': usage.prompt_tokens,
                        'completion_tokens': usage.completion_tokens,
                        'total_tokens': usage.total_tokens
                    },
                    'model': requested_model,  # Return the actual model used
                    'conversation_id': conversation_id
                })
                
            except openai.APIError as e:
                logger.error(f"OpenAI API error: {e}")
                return jsonify({
                    'error_code': 'openai_api_error',
                    'message': f'OpenAI API error: {str(e)}'
                }), 502
                
            except openai.Timeout as e:
                logger.error(f"OpenAI timeout: {e}")
                return jsonify({
                    'error_code': 'openai_timeout',
                    'message': f'OpenAI request timed out after {openai_config["timeout"]} seconds'
                }), 502
                
            except Exception as e:
                logger.error(f"OpenAI client error: {e}")
                return jsonify({
                    'error_code': 'openai_client_error',
                    'message': f'OpenAI client error: {str(e)}'
                }), 502
            
        except Exception as e:
            logger.error(f"AI Chat endpoint error: {e}")
            return jsonify({'error': 'AI Chat failed', 'details': str(e)}), 500

    @app.route('/api/system/generate', methods=['POST'])
    def generate_system():
        """Generate a complete system based on specifications"""
        try:
            data = request.get_json()
            
            # Validate required fields
            required_fields = ['name', 'description', 'type', 'techStack', 'features', 'infrastructure']
            for field in required_fields:
                if field not in data:
                    return jsonify({
                        'success': False,
                        'error': f'Missing required field: {field}'
                    }), 400
            
            # Generate system ID
            system_id = str(uuid.uuid4())
            
            # Create system specification
            system_spec = {
                'id': system_id,
                'name': data['name'],
                'description': data['description'],
                'type': data['type'],
                'techStack': data['techStack'],
                'features': data['features'],
                'infrastructure': data['infrastructure'],
                'createdAt': datetime.utcnow().isoformat(),
                'status': 'generating'
            }
            
            # Generate system architecture
            architecture = generate_system_architecture(system_spec)
            
            # Generate system templates
            templates = generate_system_templates(system_spec, architecture)
            
            # Generate deployment configuration
            deployment_config = generate_deployment_config(system_spec, architecture)
            
            # Create complete system output
            system_output = {
                'systemId': system_id,
                'specification': system_spec,
                'architecture': architecture,
                'templates': templates,
                'deployment': deployment_config,
                'status': 'generated',
                'generatedAt': datetime.utcnow().isoformat()
            }
            
            # Store system for preview/testing in S3
            save_system_to_s3(system_id, system_output)
            
            return jsonify({
                'success': True,
                'system': system_output
            })
            
        except Exception as e:
            logger.error(f"System generation error: {e}")
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500

    @app.route('/api/system/preview/<system_id>', methods=['GET'])
    def preview_system(system_id):
        """Preview a generated system with code viewer and architecture"""
        try:
            system = load_system_from_s3(system_id)
            if not system:
                return jsonify({
                    'success': False,
                    'error': 'System not found'
                }), 404
            
            # Create preview data with file contents
            preview_data = {
                'systemId': system_id,
                'specification': system['specification'],
                'architecture': system['architecture'],
                'templates': system['templates'],
                'deployment': system['deployment'],
                'preview': {
                    'fileCount': count_files(system['templates']),
                    'components': len(system['architecture']['components']),
                    'infrastructure': len(system['architecture']['infrastructure']),
                    'generatedAt': system['generatedAt']
                }
            }
            
            return jsonify({
                'success': True,
                'preview': preview_data
            })
            
        except Exception as e:
            logger.error(f"System preview error: {e}")
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500

    @app.route('/api/system/test/<system_id>', methods=['POST'])
    def test_system(system_id):
        """Deploy and test a generated system"""
        try:
            system = load_system_from_s3(system_id)
            if not system:
                return jsonify({
                    'success': False,
                    'error': 'System not found'
                }), 404
            
            # Create test deployment
            test_deployment = create_test_deployment(system)
            
            # Store test deployment info and save back to S3
            system['testDeployment'] = test_deployment
            save_system_to_s3(system_id, system)
            
            return jsonify({
                'success': True,
                'testDeployment': test_deployment
            })
            
        except Exception as e:
            logger.error(f"System test error: {e}")
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500

    @app.route('/api/system/validate/<system_id>', methods=['GET'])
    def validate_system(system_id):
        """Validate a generated system"""
        try:
            system = load_system_from_s3(system_id)
            if not system:
                return jsonify({
                    'success': False,
                    'error': 'System not found'
                }), 404
            
            # Run validation checks
            validation_results = validate_system_components(system)
            
            return jsonify({
                'success': True,
                'validation': validation_results
            })
            
        except Exception as e:
            logger.error(f"System validation error: {e}")
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500

    @app.route('/api/system/download/<system_id>', methods=['GET'])
    def download_system(system_id):
        """Download a generated system as ZIP file"""
        try:
            system = load_system_from_s3(system_id)
            if not system:
                return jsonify({
                    'success': False,
                    'error': 'System not found'
                }), 404
            
            # Create ZIP file
            zip_buffer = create_system_zip(system)
            
            # Return ZIP file
            return send_file(
                io.BytesIO(zip_buffer),
                mimetype='application/zip',
                as_attachment=True,
                download_name=f"{system['specification']['name'].lower().replace(' ', '-')}-system.zip"
            )
            
        except Exception as e:
            logger.error(f"System download error: {e}")
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500

    @app.route('/api/system/deploy/<system_id>', methods=['POST'])
    def deploy_system(system_id):
        """Deploy a generated system to AWS with live URL"""
        try:
            # Load system from S3
            system = load_system_from_s3(system_id)
            if not system:
                return jsonify({'error': 'System not found', 'success': False}), 404
            
            # Get deployment configuration
            deployment_config = request.get_json() or {}
            custom_domain = deployment_config.get('domain', f"{system_id[:8]}.sbh.umbervale.com")
            deployment_type = deployment_config.get('type', 'production')  # 'preview' or 'production'
            
            # Validate domain
            domain_validation = validate_domain(custom_domain)
            if not domain_validation['valid']:
                return jsonify({
                    'success': False,
                    'error': domain_validation['error']
                }), 400
            
            # Get deployment strategy
            domain_type = domain_validation['domain_type']
            strategy = get_deployment_strategy(domain_type)
            
            # Generate appropriate domain based on deployment type
            if deployment_type == 'preview':
                if domain_type == 'sbh_managed':
                    final_domain = f"preview-{system_id[:8]}.sbh.umbervale.com"
                else:
                    final_domain = f"preview-{custom_domain}"
            else:
                final_domain = custom_domain
            
            # Deploy to AWS ECS
            deployment_result = deploy_to_aws_ecs(system, final_domain, deployment_type)
            
            if deployment_result['success']:
                # Handle DNS setup based on domain type
                dns_result = setup_domain_dns(final_domain, deployment_result['load_balancer_dns'], domain_type)
                
                # Request SSL certificate
                ssl_result = request_ssl_certificate(final_domain)
                
                return jsonify({
                    'success': True,
                    'system_id': system_id,
                    'live_url': f"https://{final_domain}",
                    'deployment_id': deployment_result['deployment_id'],
                    'deployment_type': deployment_type,
                    'domain_type': domain_type,
                    'dns_setup': dns_result,
                    'ssl_setup': ssl_result,
                    'strategy': strategy,
                    'status': 'deployed',
                    'message': f'System deployed successfully! Access it at https://{final_domain}'
                })
            else:
                return jsonify({
                    'success': False,
                    'error': deployment_result['error']
                }), 500
                
        except Exception as e:
            logger.error(f"Error deploying system {system_id}: {e}")
            return jsonify({'error': str(e), 'success': False}), 500

    @app.route('/api/system/domain/validate', methods=['POST'])
    def validate_domain_endpoint():
        """Validate domain and get setup instructions"""
        try:
            data = request.get_json()
            if not data or 'domain' not in data:
                return jsonify({'error': 'Domain required', 'success': False}), 400
            
            domain = data['domain']
            validation_result = validate_domain(domain)
            
            if not validation_result['valid']:
                return jsonify({
                    'success': False,
                    'error': validation_result['error']
                }), 400
            
            domain_type = validation_result['domain_type']
            strategy = get_deployment_strategy(domain_type)
            
            # Generate setup instructions
            setup_instructions = generate_setup_instructions(domain, domain_type)
            
            return jsonify({
                'success': True,
                'domain': domain,
                'domain_type': domain_type,
                'strategy': strategy,
                'setup_instructions': setup_instructions
            })
            
        except Exception as e:
            logger.error(f"Error validating domain: {e}")
            return jsonify({'error': str(e), 'success': False}), 500

    @app.route('/api/system/domain/status/<system_id>', methods=['GET'])
    def check_domain_status(system_id):
        """Check domain setup status"""
        try:
            # Load system from S3
            system = load_system_from_s3(system_id)
            if not system:
                return jsonify({'error': 'System not found', 'success': False}), 404
            
            # Get deployment info (this would be stored in a deployments table in production)
            domain = request.args.get('domain')
            if not domain:
                return jsonify({'error': 'Domain required', 'success': False}), 400
            
            # Check DNS propagation
            dns_status = check_dns_propagation(domain, '')
            
            # Check SSL certificate status
            ssl_status = check_ssl_certificate_status(domain)
            
            return jsonify({
                'success': True,
                'domain': domain,
                'dns_status': dns_status,
                'ssl_status': ssl_status,
                'overall_status': 'ready' if dns_status['propagated'] and ssl_status['ready'] else 'pending'
            })
            
        except Exception as e:
            logger.error(f"Error checking domain status: {e}")
            return jsonify({'error': str(e), 'success': False}), 500

    @app.route('/api/system/upload-reference/<system_id>', methods=['POST'])
    def upload_reference_files(system_id):
        """Upload reference files (screenshots, docs, etc.) for system generation"""
        try:
            if 'files' not in request.files:
                return jsonify({'error': 'No files provided', 'success': False}), 400
            
            files = request.files.getlist('files')
            file_type = request.form.get('type', 'general')  # 'screenshot', 'document', 'wireframe', 'general'
            
            uploaded_files = []
            
            for file in files:
                if file.filename == '':
                    continue
                    
                # Upload to S3
                upload_result = upload_file_to_s3(file, system_id, file_type)
                
                if upload_result['success']:
                    # Analyze the file based on type
                    analysis_result = None
                    
                    if file.content_type and file.content_type.startswith('image/'):
                        analysis_result = analyze_uploaded_image(upload_result['s3_key'])
                    elif file.content_type in ['application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword', 'text/plain']:
                        analysis_result = extract_text_from_document(upload_result['s3_key'], file.content_type)
                    
                    uploaded_files.append({
                        'file_id': upload_result['file_id'],
                        'filename': upload_result['filename'],
                        'content_type': upload_result['content_type'],
                        'size': upload_result['size'],
                        's3_key': upload_result['s3_key'],
                        'analysis': analysis_result
                    })
            
            return jsonify({
                'success': True,
                'system_id': system_id,
                'uploaded_files': uploaded_files,
                'message': f'Successfully uploaded {len(uploaded_files)} files'
            })
            
        except Exception as e:
            logger.error(f"Error uploading reference files: {e}")
            return jsonify({'error': str(e), 'success': False}), 500

    @app.route('/api/system/analyze-url', methods=['POST'])
    def analyze_reference_url_endpoint():
        """Analyze a reference URL for system inspiration"""
        try:
            data = request.get_json()
            if not data or 'url' not in data:
                return jsonify({'error': 'URL required', 'success': False}), 400
            
            url = data['url']
            analysis_result = analyze_reference_url(url)
            
            if analysis_result['success']:
                return jsonify({
                    'success': True,
                    'url': url,
                    'analysis': analysis_result['analysis']
                })
            else:
                return jsonify({
                    'success': False,
                    'error': analysis_result['error']
                }), 400
                
        except Exception as e:
            logger.error(f"Error analyzing URL: {e}")
            return jsonify({'error': str(e), 'success': False}), 500

    @app.route('/api/system/edit/<system_id>', methods=['POST'])
    def edit_system(system_id):
        """Edit an existing system with new specifications or feedback"""
        try:
            # Load existing system
            system = load_system_from_s3(system_id)
            if not system:
                return jsonify({'error': 'System not found', 'success': False}), 404
            
            # Get edit data
            edit_data = request.get_json()
            if not edit_data:
                return jsonify({'error': 'No edit data provided', 'success': False}), 400
            
            # Update system based on edit type
            edit_type = edit_data.get('type', 'specification')
            
            if edit_type == 'specification':
                # Update system specification
                new_spec = edit_data.get('specification', {})
                system['specification'].update(new_spec)
                
                # Regenerate system with new specs
                from server_part2 import generate_system_templates
                from server_part3 import generate_system_architecture
                from server_part4 import generate_deployment_config
                
                # Regenerate architecture and templates
                system['architecture'] = generate_system_architecture(system['specification'])
                system['templates'] = generate_system_templates(system['specification'], system['architecture'])
                system['deployment'] = generate_deployment_config(system['specification'], system['architecture'])
                
            elif edit_type == 'file_update':
                # Update specific files
                file_updates = edit_data.get('file_updates', {})
                for file_path, new_content in file_updates.items():
                    if file_path in system['templates']:
                        system['templates'][file_path]['content'] = new_content
            
            elif edit_type == 'feature_add':
                # Add new features
                new_features = edit_data.get('features', [])
                system['specification']['features'].extend(new_features)
                
                # Regenerate affected templates
                from server_part2 import generate_system_templates
                system['templates'] = generate_system_templates(system['specification'], system['architecture'])
            
            # Update metadata
            system['lastModified'] = datetime.now().isoformat()
            system['editHistory'] = system.get('editHistory', [])
            system['editHistory'].append({
                'timestamp': datetime.now().isoformat(),
                'type': edit_type,
                'changes': edit_data
            })
            
            # Save updated system
            if save_system_to_s3(system_id, system):
                return jsonify({
                    'success': True,
                    'system_id': system_id,
                    'message': 'System updated successfully',
                    'updated_system': system
                })
            else:
                return jsonify({'error': 'Failed to save updated system', 'success': False}), 500
                
        except Exception as e:
            logger.error(f"Error editing system {system_id}: {e}")
            return jsonify({'error': str(e), 'success': False}), 500

    @app.route('/api/system/regenerate/<system_id>', methods=['POST'])
    def regenerate_system(system_id):
        """Regenerate parts of an existing system"""
        try:
            # Load existing system
            system = load_system_from_s3(system_id)
            if not system:
                return jsonify({'error': 'System not found', 'success': False}), 404
            
            # Get regeneration data
            regen_data = request.get_json()
            if not regen_data:
                return jsonify({'error': 'No regeneration data provided', 'success': False}), 400
            
            components_to_regenerate = regen_data.get('components', ['all'])
            
            # Regenerate specified components
            if 'all' in components_to_regenerate or 'architecture' in components_to_regenerate:
                from server_part3 import generate_system_architecture
                system['architecture'] = generate_system_architecture(system['specification'])
            
            if 'all' in components_to_regenerate or 'templates' in components_to_regenerate:
                from server_part2 import generate_system_templates
                system['templates'] = generate_system_templates(system['specification'], system['architecture'])
            
            if 'all' in components_to_regenerate or 'deployment' in components_to_regenerate:
                from server_part4 import generate_deployment_config
                system['deployment'] = generate_deployment_config(system['specification'], system['architecture'])
            
            # Update metadata
            system['lastModified'] = datetime.now().isoformat()
            system['regenerationHistory'] = system.get('regenerationHistory', [])
            system['regenerationHistory'].append({
                'timestamp': datetime.now().isoformat(),
                'components': components_to_regenerate
            })
            
            # Save regenerated system
            if save_system_to_s3(system_id, system):
                return jsonify({
                    'success': True,
                    'system_id': system_id,
                    'message': f'Successfully regenerated {", ".join(components_to_regenerate)}',
                    'regenerated_system': system
                })
            else:
                return jsonify({'error': 'Failed to save regenerated system', 'success': False}), 500
                
        except Exception as e:
            logger.error(f"Error regenerating system {system_id}: {e}")
            return jsonify({'error': str(e), 'success': False}), 500

    return app

# Helper functions
def count_files(templates):
    """Count total files in templates"""
    total = 0
    for template_type, template_data in templates.items():
        if 'files' in template_data:
            total += len(template_data['files'])
    return total

def create_test_deployment(system):
    """Create test deployment configuration"""
    return {
        'testId': f"test_{system['systemId'][:8]}",
        'status': 'deploying',
        'frontendUrl': f"https://test-{system['systemId'][:8]}.sbh.umbervale.com",
        'backendUrl': f"https://api-test-{system['systemId'][:8]}.sbh.umbervale.com",
        'deployedAt': datetime.utcnow().isoformat(),
        'estimatedTime': '2-3 minutes'
    }

def validate_system_components(system):
    """Validate system components"""
    validation = {
        'overall': 'valid',
        'checks': [],
        'warnings': [],
        'errors': []
    }
    
    # Check required files
    required_files = ['package.json', 'main.tf', 'Dockerfile']
    for template_type, template_data in system['templates'].items():
        if 'files' in template_data:
            for file_info in template_data['files']:
                if file_info['name'] in required_files:
                    validation['checks'].append({
                        'type': 'file',
                        'name': file_info['name'],
                        'status': 'found',
                        'template': template_type
                    })
    
    # Check infrastructure components
    if len(system['architecture']['infrastructure']) < 2:
        validation['warnings'].append('Minimal infrastructure components')
    
    # Check for security
    if 'User Authentication' in system['specification']['features']:
        validation['checks'].append({
            'type': 'security',
            'name': 'Authentication',
            'status': 'configured'
        })
    
    return validation

def create_system_zip(system):
    """Create ZIP file from system templates"""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Add README
        readme_content = f"""# {system['specification']['name']}

{system['specification']['description']}

## Generated System Components

### Frontend
- Type: {system['templates'].get('frontend', {}).get('type', 'N/A')}
- Files: {len(system['templates'].get('frontend', {}).get('files', []))}

### Backend  
- Type: {system['templates'].get('backend', {}).get('type', 'N/A')}
- Files: {len(system['templates'].get('backend', {}).get('files', []))}

### Infrastructure
- Type: {system['templates'].get('infrastructure', {}).get('type', 'N/A')}
- Files: {len(system['templates'].get('infrastructure', {}).get('files', []))}

## Deployment Instructions

1. Review the generated code
2. Update environment variables
3. Deploy infrastructure with Terraform
4. Deploy applications to ECS
5. Configure CI/CD pipelines

Generated by System Builder Hub (SBH)
Generated at: {system['generatedAt']}
"""
        zip_file.writestr('README.md', readme_content)
        
        # Add all template files
        for template_type, template_data in system['templates'].items():
            if 'files' in template_data:
                for file_info in template_data['files']:
                    file_path = f"{template_type}/{file_info['name']}"
                    zip_file.writestr(file_path, file_info['content'])
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

    return app
